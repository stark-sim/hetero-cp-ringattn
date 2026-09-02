from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE = Path('/Users/stark_sim/Desktop/硕士课题/开题报告')
SRC = BASE / '开题报告_新版12_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx'
OUT = BASE / '开题报告_新版16_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx'
IMG_DIR = Path('/Users/stark_sim/VSCodeProjects/hetero-cp-ringattn/output/imagegen')


def clone_run_format(para):
    if para.runs and para.runs[0]._r.rPr is not None:
        return deepcopy(para.runs[0]._r.rPr)
    return None


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = clone_run_format(para)
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def find_para(doc, marker):
    for para in doc.paragraphs:
        if marker in para.text:
            return para
    raise ValueError(f'paragraph not found: {marker}')


def make_para(doc, text, style=None, bold=False, align=None, size=None):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align is not None:
        para.alignment = align
    return para


def insert_after(anchor, element):
    anchor_node = anchor._tbl if hasattr(anchor, '_tbl') else anchor._p
    element_node = element._tbl if hasattr(element, '_tbl') else element._p
    anchor_node.addnext(element_node)
    return element


def add_caption(doc, text):
    para = make_para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_figure(doc, anchor, image_path, caption, width=6.1):
    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_before = Pt(3)
    image_para.paragraph_format.space_after = Pt(0)
    image_para.add_run().add_picture(str(image_path), width=Inches(width))
    insert_after(anchor, image_para)
    cap = add_caption(doc, caption)
    insert_after(image_para, cap)
    return cap


def set_cell(cell, text, bold=False, size=9.5):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    for margin in ('top', 'start', 'bottom', 'end'):
        pass


def add_table_after(doc, anchor, title, headers, rows, widths):
    cap = add_caption(doc, title)
    insert_after(anchor, cap)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Normal Table'
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, size=8.7)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    insert_after(cap, table)
    note = make_para(doc, '注：表中“新增机制”仅指本课题拟实现或验证的增量，不把参考框架已有能力写作创新。', size=9)
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(6)
    insert_after(table, note)
    return note


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    doc = Document(SRC)

    # Add macro-to-specific transitions without changing the established scope.
    p = find_para(doc, '本课题即源于对上述问题的思考与前期实践。')
    new = make_para(doc, '从研究对象的可拆分性出发，本课题将“异构承接”理解为一组带条件的系统选择：先识别负载阶段和状态依赖，再估计通信、显存、排队与恢复成本，最后决定采用请求级、阶段级或任务内协同。这样的定义既能容纳真实环境中的设备差异，也能避免将所有连接在同一网络中的加速卡都视为可以共同执行任意任务。')
    insert_after(p, new)
    p = find_para(doc, '这些工作的实际意义在于让异构资源以可预测的方式承接真实工作')
    new = make_para(doc, '因此，课题的结论将同时包含正向结果和负向边界：正向结果说明某类设备组合、上下文规模或阶段比例下存在可复现的净收益；负向边界说明通信、拖尾、版本滞后或恢复开销超过收益时应当拒绝协同。将两类结果同时写入研究结论，能够使后续系统设计具有明确的使用条件。')
    insert_after(p, new)

    p = find_para(doc, '本章按“服务组织—任务内协同—后训练工作流—弱连接协作”的顺序梳理国外研究')
    new = make_para(doc, '阅读本章时需要区分三种证据。第一种是已经在工业或开源系统中稳定存在的工程能力，例如请求路由、阶段分离、KV管理和异步流水线；第二种是论文中在受控环境下验证的并行机制，例如非均匀序列切分和低频参数同步；第三种是本课题尚需通过原型和实验确认的组合能力，即在明确通信和恢复约束下，让异构设备承接一部分真实工作。后文将按这一层次标记研究空缺，避免把“可借鉴”误写成“已解决”。')
    insert_after(p, new)
    p = find_para(doc, '由此，本课题的两条路线可共享的分析问题可具体表述为')
    new = make_para(doc, '这一综述还给出一个重要的节奏判断：研究不应从“如何把更多设备加入任务”开始，而应从“不加入会损失什么、加入后新增什么代价”开始。推理线先验证长上下文是否存在显存墙，再决定是否打开HCP；RL线先确认阶段边界和版本约束，再决定是否把弹性设备纳入工作流。两条路线均以能够被实验拒绝的准入条件作为起点。')
    insert_after(p, new)

    # Add a domestic-evidence subsection before the synthesis, keeping positive
    # capability claims paired with explicit communication and deployment limits.
    p = find_para(doc, '这些工作支持“生成和评估阶段适合利用异构资源”的判断')
    domestic = make_para(doc, '2.2.3  昇腾集群与国产异构软件栈的工程证据', bold=True)
    insert_after(p, domestic)
    p = domestic
    for text in [
        '昇腾相关公开成果为国内异构算力研究提供了比“设备型号对比”更具体的证据。PanGu-Σ在昇腾910集群和MindSpore上完成了万亿参数稀疏语言模型训练，通过专家计算与存储分离扩大了可用的计算与存储边界[40]；该工作说明国产NPU集群能够承载大模型训练，但其收益来自模型结构、并行方式与系统软件的协同设计，并不意味着任意异构卡都可以直接混合执行高频同步训练。',
        'CloudMatrix384进一步展示了384个昇腾910 NPU与192个鲲鹏CPU通过统一总线构成的超节点，并在推理侧采用Prefill、Decode和缓存的独立扩展[41]。该案例对本课题有两点启示：一是国产集群已经把异构资源池化和阶段解耦推进到生产级系统；二是其MoE专家并行和KV访问依赖专用高带宽全互联，反向说明普通网络下的跨厂商任务内协同必须谨慎设定边界。',
        '从软件栈看，MindSpore/MindSpeed提供了数据并行、张量并行、流水并行和混合并行等分布式能力[42]，vLLM-Ascend则通过插件方式把昇腾后端接入主流推理接口[44]。这些工程能力可作为本课题的适配基线，但“后端可运行”与“跨设备共同完成一次任务”是两个不同层次的问题：前者解决算子和运行时兼容，后者还需要设备画像、通信预算、准入和失败回退。',
        '国内异构协同的学术证据也不只来自大模型。ReDSEa在鲲鹏CPU与昇腾910组成的超算异构系统上对三角方程求解进行自动映射、负载均衡和调度[43]，证明了计算密集型任务可以通过性能模型获得异构收益；但其任务具有明确的分块结构，与LLM推理或RL后训练不同。因此，本课题将昇腾集群和ReDSEa作为“异构承接可行但依赖负载结构”的国内证据，而不是直接复用其性能结论。'
    ]:
        q = make_para(doc, text)
        insert_after(p, q)
        p = q
    framing = make_para(doc, '图4将异构现状、已有系统的承接粒度和本课题拟补足的研究空档放在同一比较轴上；昇腾与CloudMatrix384作为国内工程证据嵌入主流承接方式，而不单独代表本课题的研究对象。')
    insert_after(p, framing)
    p = framing
    state = IMG_DIR / 'heterogeneous-service-landscape-v1.png'
    if state.exists():
        cap = add_figure(doc, p, state, '图4  异构算力服务现状、主流承接方式与研究空档', width=6.0)
        p = cap
    rows = [
        ('Dynamo / vLLM / llm-d', '请求路由、Prefill/Decode分离、KV管理、后端接口', '适合请求级/阶段级承接；不解决单请求任务内HCP'),
        ('CloudMatrix384 / vLLM-Ascend', '昇腾+鲲鹏专用互联；国产后端插件与阶段解耦', '证明国产异构池化可落地；依赖专用高带宽互联与软件栈'),
        ('HetRL / Prime RL', 'rollout、奖励、评估阶段放置；异步流水与弹性设备', '适合RL阶段承接；版本、队列和恢复边界仍需组合验证'),
        ('本课题', '长上下文Prefill的HCP；RL阶段能力合同与回退', '把协同粒度、准入条件和负向边界变成可测规则'),
    ]
    note = add_table_after(doc, p, '表3  主流异构承接方式与本课题增量边界', ['代表工作/对象', '已展示能力', '对本课题的启示与限制'], rows, [1.6, 2.8, 2.0])

    # Add overview figure and a compact research-plan matrix at the beginning of Chapter 3.
    p = find_para(doc, '两条路线均从相同的分析维度出发')
    overview = IMG_DIR / 'heterogeneous-two-research-directions-overview-v4-labeled.png'
    if overview.exists():
        cap = add_figure(doc, p, overview, '图1  两条独立研究路线及其异构负载承接边界', width=6.0)
        p = cap
    rows = [
        ('长上下文Prefill', '输入长度、KV规模、显存、P2P带宽/时延、TTFT约束', 'HCP准入；capacity-aware非对称seq_chunk_len/block_size；K/V P2P ring；超时和容量变化回退', '单设备、同构CP、请求/PD路由、固定均分HCP；TTFT、尾延迟、显存峰值、网络字节数、回退率'),
        ('LLM-RL后训练阶段', '阶段输入输出、后端兼容性、版本、队列、恢复窗口', '阶段能力合同；rollout/奖励评估/数据处理异构放置；版本与样本完整性检查；节点退出回退', '固定同位、静态阶段放置、异步流水线、动态阶段承接；有效样本吞吐、样本有效率、版本滞后、恢复时间'),
    ]
    note = add_table_after(doc, p, '表1  两条研究路线的方案闭环', ['研究路线', '输入与约束', '本课题新增机制', '基线与主要评价'], rows, [1.0, 1.7, 2.1, 1.6])

    p = find_para(doc, '3.2  实施方案')
    value_note = make_para(doc, '图5不再重复研究路线和实验矩阵，而是区分参考框架已经提供的能力、尚未覆盖的具体空档，以及本课题拟增加并验证的机制。推理与LLM-RL在图中保持为上下两条独立路线。')
    insert_after(p, value_note)
    matrix = IMG_DIR / 'heterogeneous-value-boundary-v1.png'
    if matrix.exists():
        add_figure(doc, value_note, matrix, '图5  现有异构系统能力与本课题新增机制的边界对比', width=6.0)

    # Route-specific figures stay inside their own sections.
    p = find_para(doc, '本研究只把长上下文Prefill作为任务内异构协同的验证对象')
    hcp = IMG_DIR / 'heterogeneous-hcp-prefill-detail-v3.png'
    if hcp.exists():
        add_figure(doc, p, hcp, '图2  长上下文Prefill中HCP的非对称切分、K/V环传递与准入回退', width=5.9)
    p = find_para(doc, '本研究以LLM-RL后训练工作流为对象')
    rl = IMG_DIR / 'heterogeneous-llm-rl-stage-admission-v3.png'
    if rl.exists():
        add_figure(doc, p, rl, '图3  LLM-RL后训练的阶段级异构承接与弹性回退', width=5.9)

    # Add a capability boundary table in the completed-work chapter.
    p = find_para(doc, '（1）形成 hetero-cp-ringattn 开源代码库和异构通信原型')
    rows = [
        ('vLLM / Dynamo类框架', 'API、请求路由、Prefill/Decode分离、KV生命周期、常规worker执行', '不重写服务框架；增加HCP候选计划的调用接口与回退状态'),
        ('HCP原型', '非均匀seq_chunk_len、block_size、K/V P2P ring、online softmax协议基础', '补充设备/链路画像、准入代价估计、容量感知切分和服务适配验证'),
        ('HetRL / Prime RL', '阶段放置、异步rollout、训练/评估工作流组织', '组合阶段能力合同、版本/队列约束、样本完整性与动态回退'),
        ('Prime DiLoCo', '弹性设备网格、低频同步、异步检查点和恢复参考', '在LLM-RL阶段工作流中验证节点加入/退出、链路降级和恢复窗口'),
    ]
    add_table_after(doc, p, '表4  参考框架已有能力与本课题新增机制', ['参考对象', '已有能力（作为基线或接口）', '本课题新增或待验证内容'], rows, [1.3, 3.0, 2.1])

    # Make key section headings visually distinct while preserving document fonts.
    for marker in ['1.  课题来源及研究的背景和意义', '2. 国内外在该方向的研究现状及分析', '3. 主要研究内容及实施方案', '4. 预 期 达 到 的 目 标', '5. 已 完 成 的 研 究 工 作 与 进 度 安 排', '6.  为完成课题已具备和所需的条件和经费', '7.  预计研究过程中可能遇到的困难和问题，以及解决的措施', '8.  主要参考文献']:
        try:
            para = find_para(doc, marker)
            if para.runs:
                para.runs[0].bold = True
        except ValueError:
            pass

    # Extend the bibliography with the domestic evidence cited in section 2.2.
    p = find_para(doc, 'Prime Intellect. PRIME-DiLoCo: Elastic and Low-Communication Distributed Training')
    refs = [
        '[40] Ren X, Zhou P, Meng X, et al. PanGu-Σ: Towards Trillion Parameter Language Model with Sparse Heterogeneous Computing[EB/OL]. arXiv:2303.10845, 2023.',
        '[41] Zuo P, Lin H, Deng J, et al. Serving Large Language Models on Huawei CloudMatrix384[EB/OL]. arXiv:2506.12708, 2025.',
        '[42] MindSpore. Distributed Parallelism Overview[EB/OL]. https://www.mindspore.cn/tutorials/experts/en/r2.3.1/parallel/overview.html, 2025.',
        '[43] Zacharopoulos G, Bournias I, Vlacic V, et al. ReDSEa: Automated Acceleration of Triangular Solver on Supercloud Heterogeneous Systems[EB/OL]. arXiv:2305.19917, 2023.',
        '[44] vLLM Project. vLLM Ascend: Community-Maintained Hardware Plugin for Running vLLM on Ascend NPU[EB/OL]. https://github.com/vllm-project/vllm-ascend, 2025.',
    ]
    for ref in refs:
        q = make_para(doc, ref, size=9)
        insert_after(p, q)
        p = q

    doc.save(OUT)
    print(f'wrote {OUT}')


if __name__ == '__main__':
    main()
