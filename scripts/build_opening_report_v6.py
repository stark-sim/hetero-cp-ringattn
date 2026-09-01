from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path(
    "/Users/stark_sim/Desktop/硕士课题/开题报告/"
    "开题报告_新版2_沈达_面向大模型训练与推理的异构算力协同可行性理论与关键技术研究.docx"
)
OUT = SRC.with_name(
    "开题报告_新版6_沈达_面向大模型推理与强化学习后训练的异构算力协同可行性理论与关键技术研究.docx"
)
NEW_TITLE = "面向大模型推理与强化学习后训练的异构算力协同可行性理论与关键技术研究"


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = deepcopy(para.runs[0]._r.rPr) if para.runs and para.runs[0]._r.rPr is not None else None
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def replace_all(doc, old, new):
    for para in doc.paragraphs:
        if old in para.text:
            set_text(para, para.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        set_text(para, para.text.replace(old, new))


def clone_text_paragraph(template, text):
    p = deepcopy(template._p)
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    r = OxmlElement("w:r")
    if template.runs and template.runs[0]._r.rPr is not None:
        r.append(deepcopy(template.runs[0]._r.rPr))
    t = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def remove_para(para):
    para._element.getparent().remove(para._element)


def insert_after(anchor, heading_template, body_template, heading, body):
    anchor._p.addnext(clone_text_paragraph(body_template, body))
    anchor._p.addnext(clone_text_paragraph(heading_template, heading))


def replace_with_paragraphs(para, template, paragraphs):
    set_text(para, paragraphs[0])
    for text in reversed(paragraphs[1:]):
        para._p.addnext(clone_text_paragraph(template, text))


def replace_section(doc, overview, old_start, old_end, heading_template, body_template, items):
    set_text(overview, overview.text)
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if old_start in p.text)
    end = next(i for i, p in enumerate(paragraphs) if old_end in p.text)
    anchor = paragraphs[start - 1]
    # Each replaced heading has one body paragraph. Remove the final body too.
    for p in paragraphs[start : end + 2][::-1]:
        remove_para(p)
    for heading, body in reversed(items):
        anchor._p.addnext(clone_text_paragraph(body_template, body))
        anchor._p.addnext(clone_text_paragraph(heading_template, heading))


def set_toc_label(para, label):
    texts = list(para.iter(qn("w:t")))
    if not texts:
        raise ValueError("TOC paragraph has no visible text")
    texts[0].text = label


def update_toc(doc):
    toc = doc.part._element.xpath(".//w:sdt//w:p")
    if len(toc) < 24:
        raise ValueError("table of contents not found")

    set_toc_label(toc[15], "3.1.1  异构推理的多层次协同扩展")
    set_toc_label(toc[16], "3.1.2  强化学习工作负载的异构承接")
    set_toc_label(toc[20], "3.2.1  异构推理协同扩展方案")
    set_toc_label(toc[21], "3.2.2  强化学习负载弹性编排方案")

    for index in (23, 22, 18, 17):
        para = toc[index]
        para.getparent().remove(para)

    update_fields = doc.settings.element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        doc.settings.element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main():
    doc = Document(SRC)

    replace_all(doc, "面向大模型训练与推理的异构算力协同可行性理论与关键技术研究", NEW_TITLE)
    replace_all(doc, "大模型训练与推理", "大模型推理与强化学习后训练")
    replace_all(doc, "大模型训练和推理", "大模型推理与强化学习后训练")
    replace_all(doc, "大模型训练与推理任务", "大模型推理与强化学习后训练任务")
    replace_all(doc, "典型大模型训练与推理任务", "典型大模型推理与强化学习后训练任务")
    replace_all(doc, "不同大模型训练与推理任务", "不同大模型推理与强化学习后训练任务")
    replace_all(doc, "分布式大模型训练与推理系统", "大模型推理与 LLM-RL 后训练系统")
    replace_all(doc, "大模型训练与推理系统", "大模型推理与 LLM-RL 后训练系统")
    replace_all(doc, "分布式训练和其他任务", "LLM-RL 后训练工作流和其他阶段性任务")
    replace_all(doc, "分布式训练、强化学习工作流或去中心化算力协作", "LLM-RL 后训练工作流或去中心化弹性协作")
    replace_all(doc, "异构分布式训练与推理系统调研", "异构推理与 LLM-RL 后训练系统调研")
    replace_all(doc, "异构训练、异构推理", "异构推理与 LLM-RL 后训练")
    replace_all(doc, "分布式训练和去中心化强化学习方向", "LLM-RL 后训练与去中心化协同方向")
    replace_all(doc, "训练、监督微调、强化学习后训练和去中心化训练", "推理与 LLM-RL 后训练")
    replace_all(doc, "大模型训练和推理", "大模型推理与 LLM-RL 后训练")
    replace_all(doc, "大模型服务与后训练", "大模型推理服务与 LLM-RL 后训练")
    replace_all(doc, "大模型训练任务", "LLM-RL 后训练任务")
    replace_all(doc, "异构训练系统", "异构大模型系统")

    current_route_para = next(p for p in doc.paragraphs if "目前针对异构环境的主流思路" in p.text)
    set_text(
        current_route_para,
        "目前针对异构环境的主流思路，是将具有不同资源需求的请求、阶段或工作流任务分配给相匹配的设备。推理系统中出现了 Prefill 与 Decode 分离部署、面向异构资源池的请求路由等设计；强化学习工作流中则可将样本生成、奖励评估、策略更新等阶段分配给性能匹配的设备组。这些方法说明异构资源可以在服务和工作流层形成互补，但还缺少一套依据负载可拆分性、通信强度、设备能力和服务目标选择协同层次的统一方法。特别是张量并行、数据并行等高频同步训练方式对跨设备通信要求极高，不能简单假设异构设备共同承担同一计算即可获得收益；而长上下文 Prefill 等特定场景则可能在通信预算可接受时采用更细粒度的协同。",
    )

    source_para = next(p for p in doc.paragraphs if "本课题即源于" in p.text)
    set_text(
        source_para,
        "本课题即源于对上述问题的思考与前期实践。申请人围绕异构环境下基于环结构的上下文并行通信开展研究，已完成 hetero-cp-ringattn 原型系统的核心代码开发，并在小型混合集群上获得了初步加速效果。在此基础上，课题以大模型推理与 LLM-RL 后训练系统为主要对象：在推理侧参考 Dynamo 的请求编排、Prefill/Decode 分离和可插拔执行接口，探索异构资源池与 HCP 类细粒度协同模块的衔接；在强化学习侧参考 HetRL 的异构工作流调度，以及 Prime RL 和 Prime DiLoCo 在异步后训练、弹性设备网格和弱同步通信方面的实践，研究可拆分负载在异构资源上的承接与编排。",
    )

    layout_para = next(p for p in doc.paragraphs if "从研究布局来看" in p.text)
    set_text(
        layout_para,
        "从研究布局来看，现有工作呈现明显的层次割裂。Dynamo 等工业系统主要关注大模型推理服务的资源编排、Prefill/Decode 分离、KV 传输和执行引擎组织，其价值在于提供请求级和阶段级调度抽象，而非已经解决异构设备共同执行单一任务的问题；HetRL 关注包含 rollout、奖励评估、策略更新等阶段的异构 LLM-RL 后训练工作流；Prime RL 和 Prime DiLoCo 则展示了异步后训练、弹性设备网格、弱同步通信、节点动态加入退出和跨地域协作的工程可行性。这些工作分别覆盖服务层、工作流层和通信基础设施层，但尚缺少依据通信约束选择协同粒度，并把异构能力从资源池稳健地扩展到推理执行和强化学习工作流的方法。",
    )

    scale_para = next(p for p in doc.paragraphs if "规模层面的矛盾同样突出" in p.text)
    set_text(
        scale_para,
        "规模层面的矛盾同样突出。现有异构系统普遍沿用集中式调度架构，隐含全局拓扑可见、链路均匀可靠的假设；当节点规模扩大、链路质量参差不齐时，单一调度器的状态采集与决策开销会增加。Prime DiLoCo 的 ElasticDeviceMesh、心跳失效检测、动态进程组调整、异步检查点和弱同步通信，说明在 LLM-RL 后训练中可以把节点弹性和弱连接协作作为工程扩展；Prime RL 则将异步 rollout、训练、评估、FSDP/EP/CP 和 vLLM 执行引擎整合到统一工作流中。因此，去中心化并非本课题的既定前提，而是在跨域资源发现、局部故障恢复或弱连接协作成为瓶颈时可选的支撑机制；核心仍是针对设备能力和通信条件建立可解释的协同决策方法。",
    )

    value_para = next(p for p in doc.paragraphs if "就研究价值而言" in p.text)
    set_text(
        value_para,
        "就研究价值而言，本课题首先希望建立适用于大模型推理和 LLM-RL 后训练的异构算力协同可行性判定模型，为给定的设备组合、网络条件和任务图选择请求级、阶段级或任务内协同的合适粒度，并给出资源分配与效率边界。其次，课题将在技术层面打通从建模到系统的通路：在推理场景中探索把 HCP 的容量感知非均匀协同作为可插拔模块接入 vLLM 或 Dynamo 类执行框架，在通信预算可接受的长上下文 Prefill 中进行细粒度验证；在 LLM-RL 后训练场景中研究 rollout、评估和策略更新等可拆分阶段的异构承接、异步流水线和弹性编排。",
    )

    review_para = next(p for p in doc.paragraphs if "从上述进展可以看出" in p.text)
    set_text(
        review_para,
        "从上述进展可以看出，异构协同的关键不在于强行让所有设备共同执行高通信耦合的单一任务，而在于依据负载可拆分性和通信约束选择协同粒度：资源池和请求级调度适合大多数在线推理负载，阶段级编排适合强化学习后训练工作流，只有在长上下文 Prefill 等通信预算可接受的场景中，才进一步探索 HCP 类任务内协同。去中心化拓扑可为跨域资源发现、容错和弱连接协作提供补充，但不作为每类负载都必须采用的机制。",
    )

    direction_para = next(p for p in doc.paragraphs if "基于以上分析，本课题拟以" in p.text)
    set_text(
        direction_para,
        "基于以上分析，本课题拟以异构负载的分层协同扩展为主线，从可行性建模入手，分别在推理服务和强化学习工作流中验证请求级、阶段级与条件性任务内协同的适用边界，并在确有需要时引入弱中心或去中心化支撑机制，形成面向异构大模型基础设施的可验证技术方案。",
    )

    # 3.1: five small points -> two load-oriented research points.
    overview = next(p for p in doc.paragraphs if "本课题以异构算力协同的可行性" in p.text)
    set_text(
        overview,
        "本课题围绕异构算力协同的可行性与可持续性，研究对象限定为两类具有代表性的负载：一是大模型在线推理，研究从异构资源池、请求路由到条件性任务内协同的扩展路径；二是强化学习相关工作负载，以 LLM-RL 后训练为主要实例，研究 rollout、奖励评估和策略更新等可拆分阶段的异构承接。两点共享设备—网络—任务可行性建模和容量感知调度方法，但不预设所有异构设备均需共同执行一个高通信耦合任务。去中心化和多智能体强化学习仅在集中式或分层基线无法满足跨域资源发现、容错或动态决策需求时作为可选扩展。",
    )
    h_template = next(p for p in doc.paragraphs if "3.1.1" in p.text)
    b_template = next(p for p in doc.paragraphs if "研究面向异构算力协同的统一性能" in p.text)
    items = [
        (
            "3.1.1  面向异构推理的多层次协同扩展",
            "面向大模型在线推理研究异构能力从资源池到执行层的渐进式扩展。首先刻画设备计算速率、显存容量、链路带宽、启动开销和可靠性，建立请求级、阶段级与任务内协同的可行性和代价模型；其次参考 Dynamo 的 Prefill/Decode 分离、请求编排、KV 传输和可插拔执行接口，研究异构设备上的请求路由、阶段放置、资源配比和计算—通信联合调度；最后以长上下文 Prefill 为条件性验证场景，探索将 hetero-cp-ringattn 的容量感知非均匀切分、环序和 KV 块传输封装为可接入 vLLM 或 Dynamo 类框架的协同模块。HCP 只在通信预算与正确性约束满足时用于任务内协同，不将异构设备共同完成单一推理任务作为普遍前提。",
        ),
        (
            "3.1.2  面向强化学习工作负载的异构承接与弹性编排",
            "面向强化学习相关工作负载研究可拆分阶段在异构资源上的承接、编排与弹性执行，以 LLM-RL 后训练为主要验证实例。围绕 rollout、奖励计算、评估、策略更新和检查点等阶段，研究设备能力感知的任务放置、依赖编排、并发度调节和结果/参数传递；参考 HetRL 的阶段级异构调度，结合 Prime RL 的异步执行以及 Prime DiLoCo 的弹性设备网格、弱同步通信、节点动态加入退出和异步检查点，研究弱连接环境下的吞吐、容错与资源利用率。去中心化仅作为跨域资源发现、局部故障恢复或弱中心协作的可选扩展；多智能体强化学习也仅作为在线调度候选方法，需在启发式与约束优化基线不足时再引入，且不与被调度的 LLM-RL 负载混同。",
        ),
    ]
    replace_section(doc, overview, "3.1.1", "3.1.5", h_template, b_template, items)

    # 3.2: five implementation schemes -> two corresponding schemes.
    scheme = next(p for p in doc.paragraphs if p.text.strip() == "3.2  研究方案")
    h_template = next(p for p in doc.paragraphs if "3.2.1" in p.text)
    b_template = next(p for p in doc.paragraphs if "首先通过微基准测试建立设备" in p.text)
    items = [
        (
            "3.2.1  异构推理协同扩展方案",
            "首先通过微基准测试建立设备与链路性能档案，将异构资源池抽象为带属性的图 G=(V,E)，并把请求、Prefill、Decode、KV 传输和同步关系表示为任务图。其次以集中式或分层调度为基线：上层完成候选设备、SLO 与成本感知的请求路由，中层决定 Prefill/Decode 阶段放置、资源配比和流水化参数，底层通过可插拔后端执行模型服务与 KV 传输。对于长上下文且通信预算允许的请求，再调用 HCP 协同模块决定非均匀序列分片、环序和 KV 块传输；验证其正确性、计算通信重叠和 P2P 回退路径。对于跨域节点失效或链路抖动，优先采用局部状态上报与回退路由，必要时再研究局部邻居视图和逻辑环重构。实验比较单设备、静态异构资源池、动态请求调度与条件性 HCP 扩展，报告 TTFT、吞吐、显存峰值、通信开销和恢复时间。",
        ),
        (
            "3.2.2  强化学习工作负载的异构承接与弹性编排方案",
            "将 LLM-RL 后训练表示为包含 rollout、奖励/评估、策略更新和检查点同步的有向任务图，结合设备画像、网络状态和阶段依赖建立资源分配与可行性约束。参考 HetRL 的异构阶段级调度方式，研究不同设备组在各阶段的放置、并发和结果/参数传递；参考 Prime RL 的异步 rollout—训练—评估流水线，研究在集中式或分层协调下的稳定基线。对于跨域节点加入退出、局部失效和弱连接等情形，再参考 Prime DiLoCo 的 ElasticDeviceMesh、心跳检测、弱同步通信和异步检查点，评估弱中心或 P2P 扩展的收益与代价。动态决策先采用启发式与约束优化，只有当局部观测和状态变化使其无法有效决策时，才引入多智能体强化学习并保留安全回退策略。实验重点比较阶段级异构编排、静态同构基线、动态调度与条件性弹性协同在吞吐、策略更新延迟、资源利用率、通信量和恢复时间上的差异。",
        ),
    ]
    replace_section(doc, scheme, "3.2.1", "3.2.5", h_template, b_template, items)

    # Objectives and schedule should expose the same two-load boundary.
    for p in doc.paragraphs:
        if "建立设备—网络—任务统一的异构算力协同模型" in p.text:
            set_text(p, "（1）理论目标：建立设备—网络—任务统一的异构算力协同模型，给出面向大模型推理与 LLM-RL 后训练的协同粒度选择、可行性判定、资源约束表达和效率边界分析。")
        elif "形成容量感知的非均匀任务划分" in p.text:
            set_text(p, "（2）方法目标：形成面向推理与 LLM-RL 后训练的容量感知路由、阶段编排、条件性任务内协同、计算—通信联合调度和局部容错回退方法，并明确集中式、弱中心和多智能体决策的适用边界。")
        elif "并为分布式训练、强化学习工作流和其他任务保留统一接口" in p.text:
            set_text(p, "（3）系统目标：实现可插拔的异构协同原型系统，以 hetero-cp-ringattn 为长上下文推理的条件性协同载体，探索其与 vLLM 或 Dynamo 类框架的衔接，并为 LLM-RL 后训练工作流保留阶段级异构承接接口。")
        elif "根据前期结果选择" in p.text and "扩展验证场景" in p.text:
            set_text(p, "（4）实验目标：在多代、多厂商或不同设备能力的混合环境中，对比单设备、同构基线、静态异构编排、动态调度和条件性协同方法，系统评估推理与 LLM-RL 后训练的完成时间、吞吐、显存、通信、扩展性和故障恢复能力。")
        elif "不预设单一固定加速倍数" in p.text:
            set_text(p, "（4）实验目标：在多代、多厂商或不同设备能力的混合环境中，对比单设备、同构基线、静态异构编排、动态调度和条件性协同方法，系统评估推理与 LLM-RL 后训练的完成时间、吞吐、显存、通信、扩展性和故障恢复能力；不预设单一固定加速倍数，而是报告不同条件下的收益区间和失效边界。")
        elif "围绕统一可行性建模" in p.text:
            set_text(p, "（5）成果目标：围绕协同粒度可行性建模、异构推理扩展、强化学习工作流编排和条件性弹性机制形成可拆分的论文与专利成果，持续维护并开源可复用的系统代码和实验基准。")
        elif "已完成对异构推理与" in p.text:
            set_text(p, "已完成对异构推理与 LLM-RL 后训练、上下文并行、Prefill/Decode 分离及弹性协同系统的调研，重点分析异构请求路由、阶段级承接、P2P 协同和多智能体动态调度的适用条件，初步明确请求级、阶段级与条件性任务内协同的研究空间。")
        elif "异构设备画像与去中心化协同探索" in p.text:
            set_text(p, "5.1.1.3 异构设备画像与弹性协同探索")
        elif "已围绕设备算力、显存容量、链路带宽和通信延迟" in p.text:
            set_text(p, "已围绕设备算力、显存容量、链路带宽和通信延迟开展初步测量，完成动态环构建、非均匀切片和局部拓扑变化的仿真探索，并保留跨 CUDA、ROCm 与 MPS 后端协同验证所需的接口设计。后续将把这些工程结果统一纳入协同粒度可行性模型，并以集中式或分层调度为基线评估弹性扩展的必要性。")
        elif "形成 hetero-cp-ringattn 开源代码库" in p.text:
            set_text(p, "（1）形成 hetero-cp-ringattn 开源代码库和异构通信原型；（2）完成异构推理、LLM-RL 后训练、上下文并行与弹性协同方向的文献和系统调研；（3）搭建异构协同仿真与实验环境，具备开展设备画像、通信性能、条件性任务内协同和端到端推理验证的条件；（4）围绕异构请求调度、通信机制和工作流承接形成阶段性技术报告，并将根据后续实验结果拆分论文和专利选题。")
        elif "去中心化环境下通信环的频繁重构" in p.text:
            set_text(p, "在跨域或弱连接条件下，局部重构和 P2P 扩展可能引发路由抖动，进而影响请求或工作流的时延稳定性。对此，课题将默认采用集中式或分层调度与回退路由；只有基线在资源发现、故障恢复或弱连接协作上不足时，才采用局部一致性协议、版本化握手和短暂双轨过渡等机制，并单独评估其收益与开销。")
        elif "去中心化多智能体强化学习面对的状态空间" in p.text:
            set_text(p, "多智能体强化学习若用于异构调度，可能面对状态空间庞大、奖励信号稀疏和训练初期收敛缓慢的问题。因此本课题不将其作为既定实现路线，而是先以启发式和约束优化建立可解释的稳定基线；仅当局部观测和动态变化使基线无法满足需求时，再在小规模场景采用课程学习和集中式训练、分散式执行进行对照验证，并保留回退策略。")

    for row in doc.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "根据前期结果选择" in p.text:
                    set_text(p, "根据前期结果完成 LLM-RL 后训练工作流的扩展验证；仅在跨域资源发现或弱连接协作确有收益时，补充去中心化弹性协作实验。")
                elif "研究多智能体强化学习动态适应" in p.text:
                    set_text(p, "完成异构推理的端到端验证；比较静态编排、动态调度与条件性 HCP 扩展。若基线暴露出局部观测决策瓶颈，再开展多智能体强化学习对照实验。")
                elif "研究容量感知任务划分、计算" in p.text:
                    set_text(p, "研究容量感知请求路由、阶段放置、计算—通信联合调度和 P2P 回退机制；完成 vLLM 或 Dynamo 类框架衔接的原型设计。")

    # Version 5: derive the research gap and the research plan from both positive
    # evidence and explicit feasibility boundaries rather than from vision alone.
    body_template = next(p for p in doc.paragraphs if "产业界的动向为这一判断" in p.text)
    summary_start = next(p for p in doc.paragraphs if "产业界的动向为这一判断" in p.text)
    summary_paragraphs = [
        "从资源供给和软件形态看，异构已是大模型基础设施的客观约束，而不是单一厂商生态之外的边缘情形。新旧GPU、GPU与NPU以及不同显存和互联层级会在同一资源池中长期共存；NVLink Fusion、UALink和UEC等开放互连演进，以及国内多类超节点实践，为设备共处同一通信域创造了条件[21-26]。但“可以互联”不等价于“可以共同执行任意计算”：不同负载的计算密度、状态传递频率和时延目标不同，协同机制必须随之分层。",
        "在推理侧，现有证据直接支持请求级与阶段级的异构扩展。DistServe将Prefill与Decode解耦，并按照带宽、TTFT和TPOT约束共同决定资源配置与并行策略，表明阶段分离可以减少两类计算相互干扰[1]；Mooncake以KV Cache为中心组织独立资源池，说明长上下文服务可通过缓存与调度协同获得收益[3]；Dynamo则把KV感知路由、PD分离、SLO规划和多后端接入组织为服务层能力[4]。这些系统共同说明，异构推理首先适合把“请求到哪个资源池、哪个阶段在哪类设备上执行”作为决策对象，而非预设所有设备同步参与一次前向计算。",
        "与此同时，推理方向的反对证据也决定了研究边界。PD分离会引入KV状态迁移，收益取决于链路带宽、远端缓存访问和排队时延；DistServe本身也需要根据集群带宽放置两个阶段[1]。面向异构多阶段推理的MIST将跨集群通信时延、批处理效率和内存带宽竞争纳入分析，表明配置空间不能只按峰值算力排序[31]。通信密集型专家并行或张量并行通常依赖高带宽互连，CloudMatrix等系统正是通过全对等高带宽网络来支撑此类操作[24]。因此，本课题不把HCP式任务内协同表述为普遍优解，而把它限定为长上下文Prefill、显存压力显著且通信预算满足时的可检验扩展。",
        "在强化学习相关工作负载侧，支持依据来自阶段的天然异质性。RLHF/LLM-RL工作流同时包含生成、奖励或价值评估、训练更新、评测和检查点等任务，固定的同位部署会把不同计算和显存特征强行绑定，造成生成瓶颈与资源闲置；已有自适应放置研究通过交错或分解式部署降低冗余并提高吞吐[32]。HetRL进一步把异构GPU和网络上的RL调度建模为受约束联合优化问题，并以启发式和ILP两种方法处理解质量与求解开销的权衡[6]。Prime RL的异步rollout—训练—评估组织也说明，后训练适合从工作流粒度承接异构设备[12]。",
        "但RL工作流的可拆分性不消除一致性和弹性成本。异步rollout与策略更新之间存在数据和参数陈旧性风险；奖励、评估与训练阶段的吞吐不匹配会形成队列积压；节点加入退出还会增加检查点、状态恢复和进程组重建开销。DiLoCo能够在降低通信频率的前提下适应弱连接资源，其成立条件是以大量本地步和低频外层同步替代每步梯度交换，而不是证明常规TP/DP可直接跨弱异构链路运行[15]。Prime/INTELLECT-1与PCCL提供了弹性设备网格、异步恢复和容错通信的工程证据[14,16]，同时也把网络波动、恢复时窗和状态一致性暴露为必须度量的代价。",
        "据此，现有研究的共同缺口不是缺少某一种单点调度算法，而是缺少一套由负载可拆分性、通信强度、设备能力、服务目标与恢复代价共同决定协同粒度的可行性方法。推理侧需要回答何时停留在资源池/请求级调度、何时可升级到条件性任务内协同；RL侧需要回答何时阶段级异构承接优于固定部署、何时异步或弱中心机制的收益覆盖其一致性代价。第3章据此设置两条研究线，并将“无收益或违反约束”作为与性能提升同等重要的实验结论。",
    ]
    replace_with_paragraphs(summary_start, body_template, summary_paragraphs)
    for old_text in (
        "综合国内外研究现状可以发现",
        "在可扩展性方面，集中式调度构成了当前异构系统的共同上限",
        "通信与计算的脱节是第三个薄弱环节",
        "基于以上分析，本课题拟以异构负载的分层协同扩展为主线",
    ):
        remove_para(next(p for p in doc.paragraphs if old_text in p.text))

    overview = next(p for p in doc.paragraphs if "本课题围绕异构算力协同的可行性与可持续性" in p.text)
    overview_paragraphs = [
        "本课题围绕异构算力协同的可行性与可持续性，研究对象限定为两类具有代表性的负载：一是大模型在线推理，研究从异构资源池、请求路由到条件性任务内协同的扩展路径；二是强化学习相关工作负载，以LLM-RL后训练为主要实例，研究rollout、奖励评估和策略更新等可拆分阶段的异构承接。两点共享设备—网络—任务可行性建模和容量感知调度方法，但不预设所有异构设备均需共同执行一个高通信耦合任务。",
        "本课题将“协同粒度选择”作为贯穿两点的核心问题：先以设备计算速率、显存/缓存容量、链路带宽与时延、负载阶段依赖和SLO为输入，判断资源池级、请求级、阶段级或任务内协同中哪一层具备正的净收益；随后在满足准入条件的层级中优化放置、并发度、切分比例和通信路径。若估计收益不足以覆盖状态迁移、同步或恢复开销，系统应回退到更粗粒度的编排，而不是为了形式上的异构合作强行建立跨设备耦合。",
        "研究的评价不仅报告平均吞吐提升，也同时报告尾延迟、显存峰值、KV或参数传输量、调度/重构开销、故障恢复时间和策略质量变化。这样，正结果给出可复用的适用区间，负结果则界定不可用的网络、上下文长度、负载比例或异构度范围；二者共同构成面向后续系统扩展的可行性理论。去中心化与多智能体强化学习不作为预设研究对象，而分别作为跨域弹性和在线决策的条件性候选机制。",
    ]
    replace_with_paragraphs(overview, body_template, overview_paragraphs)

    inference_content = next(p for p in doc.paragraphs if "面向大模型在线推理研究异构能力" in p.text)
    inference_paragraphs = [
        "面向大模型在线推理研究异构能力从资源池到执行层的渐进式扩展。首先刻画设备计算速率、显存容量、链路带宽、启动开销和可靠性，建立请求级、阶段级与任务内协同的可行性和代价模型。模型以TTFT、TPOT、吞吐和单位请求资源占用为服务约束，同时显式计入KV传输、排队、批处理重组与路由切换开销；其输出不是单一的“最优设备”，而是给出推荐协同粒度及可行的候选资源集合。",
        "在请求级和阶段级，参考Dynamo的Prefill/Decode分离、请求编排、KV传输和可插拔执行接口，研究异构设备上的请求路由、阶段放置、资源配比和计算—通信联合调度。重点不是仅用设备峰值算力排序，而是结合输入长度、预期输出长度、KV命中率、队列状态和链路条件，避免把短请求、长上下文请求和高KV复用请求混入同一静态规则。对于预估违反SLO或远端状态迁移代价过高的请求，调度器应保留本地执行、延迟准入或重路由等回退选择。",
        "在任务内层，以长上下文Prefill为条件性验证场景，探索将hetero-cp-ringattn的容量感知非均匀切分、环序和KV块传输封装为可接入vLLM或Dynamo类框架的协同模块。模块的准入条件包括：单设备显存或TTFT成为主要瓶颈、切分后各设备仍有正的计算贡献、预测环传输时间不主导端到端时延、以及数值正确性满足预设阈值。若任一条件不满足，系统不进入HCP路径，而保留请求/阶段级编排。该设计把HCP定位为可替换、可证伪的细粒度能力，而非要求异构设备共同完成单一推理任务的固定架构。",
    ]
    replace_with_paragraphs(inference_content, body_template, inference_paragraphs)

    rl_content = next(p for p in doc.paragraphs if "面向强化学习相关工作负载研究可拆分阶段" in p.text)
    rl_paragraphs = [
        "面向强化学习相关工作负载研究可拆分阶段在异构资源上的承接、编排与弹性执行，以LLM-RL后训练为主要验证实例。将rollout、奖励计算、评估、策略更新和检查点表示为带数据依赖、参数版本和资源需求的有向任务图，分别刻画其计算密度、显存需求、输出数据量、可容忍的版本滞后和恢复代价。由此识别哪些阶段适合部署在高吞吐推理设备、哪些阶段需要高显存训练设备，以及哪些阶段可以异步并发或批量提交。",
        "在稳定集群中，研究设备能力感知的任务放置、依赖编排、并发度调节和结果/参数传递。参考HetRL的阶段级异构调度，先以启发式与约束优化构建可解释的基线：例如以单位rollout成本、奖励评估时延、更新步吞吐和队列长度为约束，优化阶段设备组规模和流水线深度；再与固定同位部署、静态异构放置进行对照。研究重点是使工作流中最慢阶段不再由单一设备类别长期主导，而不是将高频梯度通信强行跨弱链路分散。",
        "在节点动态加入退出、跨域链路波动或检查点成本显著时，参考Prime RL、Prime DiLoCo和PCCL的异步执行、弹性设备网格、弱同步通信和状态恢复思路，评估弱中心或P2P扩展的收益与代价。扩展仅在其减少停机时间、提高有效资源利用率或扩大可用资源池时启用；若版本滞后导致策略质量下降、恢复时窗阻塞更新或通信开销超过吞吐收益，则回退到集中式或分层协调。多智能体强化学习同样仅作为在线调度候选方法，需在启发式与约束优化难以处理局部观测和非平稳状态时再引入，且不与被调度的LLM-RL负载混同。",
    ]
    replace_with_paragraphs(rl_content, body_template, rl_paragraphs)

    inference_plan = next(p for p in doc.paragraphs if "首先通过微基准测试建立设备与链路性能档案" in p.text)
    inference_plan_paragraphs = [
        "首先通过微基准测试建立设备与链路性能档案。画像至少包含各设备在不同batch和上下文长度下的Prefill/Decode吞吐、显存可用容量、KV块读写与传输速率、启动时间、P2P带宽与时延及其波动范围。将异构资源池抽象为带属性的图G=(V,E)，把请求、Prefill、Decode、KV传输和同步关系表示为任务图；在离线阶段拟合阶段时延与资源占用模型，在在线阶段用滑动窗口修正模型误差。",
        "其次以集中式或分层调度为基线。上层按照SLO、输入/输出长度预测、KV复用机会和资源价格选择候选资源池；中层决定Prefill/Decode阶段放置、并发度、批处理与流水化参数；底层通过可插拔后端执行模型服务和KV传输。基线必须包含单设备、同构池、静态异构路由和动态异构路由四种对照，以区分“异构设备本身带来的容量增益”与“调度策略带来的增益”。对于SLO风险高或状态迁移代价过大的请求，明确采用本地执行、重试或拒绝准入等安全回退，而不以平均吞吐掩盖尾部损失。",
        "对于长上下文且通信预算允许的请求，再进入HCP协同模块。模块根据设备容量和测得的计算/通信比生成非均匀序列分片、候选环序及KV块大小，并先以代价模型筛除预计通信主导或小设备拖尾严重的配置；通过版本化握手、P2P回退和局部重构保障执行连续性。正确性验证以同模型单设备或同构并行输出为对照，性能验证同时报告TTFT、TPOT、吞吐、显存峰值、网络字节数、调度决策时间和恢复时间。若任务内协同未超过请求级动态调度，或违反任一SLO/正确性约束，即将其记录为适用边界而非继续扩大部署。",
        "最后在不同异构度、上下文长度、并发度和网络条件下构造系统性实验矩阵。研究将分析HCP准入阈值随显存差异、带宽/时延和KV规模的变化，比较HCP与PD分离、仅KV感知路由及静态上下文并行的收益区间。预期产出包括可复现实验基准、协同粒度选择规则和面向vLLM或Dynamo类框架的接口原型，而不承诺在所有异构组合上获得固定加速倍数。",
    ]
    replace_with_paragraphs(inference_plan, body_template, inference_plan_paragraphs)

    rl_plan = next(p for p in doc.paragraphs if "将 LLM-RL 后训练表示为包含 rollout" in p.text)
    rl_plan_paragraphs = [
        "将LLM-RL后训练表示为包含rollout、奖励/评估、策略更新和检查点同步的有向任务图，并为每个节点记录模型副本、算力/显存需求、输入输出规模、参数版本和可容忍等待时间。结合设备画像、网络状态和阶段依赖建立资源分配与可行性约束：rollout与评估关注生成吞吐和队列延迟，策略更新关注训练显存和有效token吞吐，检查点与参数分发关注传输量和恢复窗口。该抽象允许用同一模型比较固定部署、阶段级异构放置及异步流水线，而不混淆服务调度与优化算法本身。",
        "在稳定资源池上，先实现集中式或分层协调下的阶段级异构编排基线。参考HetRL的思想，以设备能力、任务依赖和网络约束决定不同设备组在各阶段的放置、并发和结果/参数传递；以启发式与ILP/约束优化分别提供低开销在线决策和小规模近似最优对照。随后接入Prime RL式异步rollout—训练—评估流水线，设置队列长度、版本滞后和更新间隔的上限，并比较其与同步流水线在后训练吞吐、策略更新延迟、样本有效性及资源利用率上的差异。",
        "在弹性扩展阶段，模拟或实测节点加入退出、链路带宽下降和检查点恢复等事件。参考Prime DiLoCo的ElasticDeviceMesh、心跳检测、异步检查点和弱同步机制，分别测量重构时间、失效期间的有效吞吐、恢复后的参数版本一致性和策略质量变化。去中心化或P2P机制只作为对照路线：若其在资源发现、局部故障恢复或跨域协作上未显著优于集中式/分层基线，便不将其纳入主路径。这样可避免把“可加入更多节点”误写为“任何节点加入都提高有效训练效率”。",
        "动态调度实验先采用规则、启发式和约束优化；仅在工作负载和网络状态呈现明显局部观测、强非平稳且基线频繁违反约束时，再构造多智能体强化学习对照。所有方法均使用相同的资源预算、模型版本与任务集合，并报告端到端吞吐、单位有效样本成本、版本滞后、通信量、故障恢复时间以及任务质量指标。若动态方法的决策开销、探索成本或不稳定性超过收益，应以基线为最终方案，并将该结果作为方法适用边界。",
    ]
    replace_with_paragraphs(rl_plan, body_template, rl_plan_paragraphs)

    # Version 6: supplement the v5 narrative with independently retrieved
    # evidence organized by load type, evidence strength, and failure boundary.
    # The goal is to keep the two-point structure while making the argument
    # falsifiable rather than dependent on the user's named systems alone.
    inference_evidence = next(
        p for p in doc.paragraphs if "面向大模型在线推理研究异构能力从资源池到执行层" in p.text
    )
    set_text(
        inference_evidence,
        "面向大模型在线推理，独立研究已经覆盖资源池、阶段流水线和任务内分片三个层次。LLM-PQ在异构GPU集群上联合优化阶段感知模型划分、自适应量化与微批大小，并在11类生产集群上报告吞吐提升；FastDecode则把解码中的KV访问与模型计算拆成不同特征的流水段，利用跨节点CPU资源缓解GPU显存压力。这些结果属于直接系统实验，支持“异构资源可通过角色分工和阶段划分产生净收益”，但收益依赖模型质量约束、内存层级和传输路径，不能外推为任意厂商GPU都适合任务内并行。",
    )
    inference_evidence2 = next(
        p for p in doc.paragraphs if "在请求级和阶段级，参考Dynamo的" in p.text
    )
    set_text(
        inference_evidence2,
        "在请求级和阶段级，除Dynamo等工业框架外，GoodServe从异构资源上的agentic inference出发，以输出长度和GPU状态预测、SLO风险校正及运行时迁移优化goodput；MIST则通过真实硬件轨迹与分析模型刻画RAG、KV检索、Prefill和Decode在多级内存及跨集群通信下的联合作用；FlowKV进一步表明KV传输本身可能成为阶段分离的决定性开销。它们共同支持把输入/输出长度、KV命中、队列状态和链路条件纳入路由，但也反向说明平均吞吐不能替代尾延迟、迁移代价和缓存一致性的评估。",
    )
    inference_evidence3 = next(
        p for p in doc.paragraphs if "在任务内层，以长上下文Prefill为条件性验证场景" in p.text
    )
    set_text(
        inference_evidence3,
        "在任务内层，HexiSeq等长上下文研究说明非均匀序列/头分配在特定混合GPU和长序列条件下具有可测收益，但其主要证据仍集中于训练或受控集群。CloudMatrix384的反向证据尤其重要：通信密集的MoE专家并行依赖专门的高带宽全互联，而不是普通网络上的异构卡简单拼接。因此，本课题把HCP扩展限定为显存墙或Prefill瓶颈明确、P2P测量达到准入阈值、计算通信可重叠且正确性通过的请求；否则回退到资源池或阶段级调度，并把“无净收益”作为有效结论。",
    )

    rl_evidence = next(
        p for p in doc.paragraphs if "面向强化学习相关工作负载研究可拆分阶段在异构资源上" in p.text
    )
    set_text(
        rl_evidence,
        "面向强化学习相关工作负载，独立证据首先来自RLHF/LLM-RL工作流自身的阶段异质性。Xiao等人的自适应放置研究将策略、奖励、价值和参考模型从固定同位部署中拆开，以交错或分解式放置减少冗余并缓解生成瓶颈；该结果直接支持把rollout、奖励评估和策略更新视作不同资源需求的任务，而不是把“RL训练”当作单一同构作业。HetRL进一步在异构GPU和网络上联合优化阶段放置，说明设备画像、任务依赖和网络约束需要同时进入调度模型。",
    )
    rl_evidence2 = next(
        p for p in doc.paragraphs if "在稳定集群中，研究设备能力感知的任务放置" in p.text
    )
    set_text(
        rl_evidence2,
        "在稳定资源池中，Prime RL和相关开源工作流提供了异步rollout—训练—评估的工程先例；但RLBoost给出的更细结论是，rollout具有近似无状态、易并行的特征，适合承接抢占式或碎片化资源，而策略训练仍需要紧耦合GPU和全互联。该证据支持把异构资源优先用于rollout、评估和数据处理，并对训练更新保留同构或局部同构约束；若将弱链路直接用于高频参数/梯度同步，通信成本可能抵消异构资源带来的吞吐。",
    )
    rl_evidence3 = next(
        p for p in doc.paragraphs if "在节点动态加入退出、跨域链路波动或检查点成本显著时" in p.text
    )
    set_text(
        rl_evidence3,
        "在动态资源环境中，DiLoCo与INTELLECT-1证明低频外层同步、弹性设备网格和异步检查点可以把跨地域资源纳入训练过程，但这属于改变同步算法和恢复协议后的可行性，不是常规TP/DP跨弱链路的直接证明。近期RolloutPipe的结果则显示，分组流水化可以在保持on-policy约束的同时提前启动训练，但其收益建立在完整样本组、前沿组调度和固定权重窗口等条件上。由此，本课题将版本滞后、样本有效性、恢复时间和策略质量列为硬约束；异步或去中心化机制只有在覆盖这些代价后才成立。",
    )

    # Replace the v5 synthesis with an explicit evidence ledger in prose.
    synthesis = next(p for p in doc.paragraphs if p.text.startswith("从资源供给和软件形态看"))
    synthesis_paragraphs = [
        "从证据类型看，异构推理和异构RL并非同一问题的两种命名。推理侧已有多篇系统工作直接在真实或仿真异构资源上测量请求路由、阶段分离、KV迁移和模型划分；RL侧的直接证据则集中于工作流阶段放置、rollout弹性和低频同步。两条证据线都支持“异构资源应按负载特征分工”，但都没有支持“异构设备无条件共同执行任意高通信计算”这一更强命题。",
        "对于异构推理，支持证据可分为三层。第一层是LLM-PQ、FastDecode等直接实验，说明相位感知划分、KV访问分工和混合资源池能够改善吞吐或容量利用；第二层是GoodServe、MIST、FlowKV等调度与建模工作，说明输出长度、缓存命中、内存层级、链路时延和排队状态会改变最优决策；第三层是把HCP用于长上下文Prefill的研究假设，仍需在目标设备与网络上实测。反对证据来自KV迁移、尾延迟、批处理重组和高带宽互联依赖，因此本课题必须把任务内协同设为准入式而非默认路径。",
        "对于异构RL，支持证据来自自适应RLHF放置、HetRL的约束联合优化、RLBoost对抢占式rollout资源的利用以及RolloutPipe对分组流水化的探索，说明工作流级解耦和资源弹性具有明确研究价值。反对证据同样明确：策略训练阶段通常需要紧耦合通信，异步执行会引入版本滞后和样本陈旧，节点动态变化还会触发检查点、进程组和恢复协议开销。因此，异构RL的主问题应是阶段承接和弹性编排，去中心化与多智能体方法只能在集中式/分层基线暴露出局部观测或恢复瓶颈时进入对照。",
        "综合来看，现有研究的共同缺口可以客观表述为“协同粒度选择缺少跨负载的统一判据”，而不是“缺少一个更大的异构框架”。本课题将设备能力、通信强度、状态传递频率、服务SLO、版本一致性和恢复代价纳入同一可行性模型，分别验证推理中的资源池/请求级/阶段级/条件性任务内协同，以及RL中的rollout/评估/更新阶段承接。每条路线都设置正向收益指标和失败退出条件，避免把单篇论文或单一系统的结果写成普遍规律。",
    ]
    replace_with_paragraphs(synthesis, body_template, synthesis_paragraphs)
    for old in (
        "在推理侧，现有证据直接支持请求级与阶段级的异构扩展",
        "与此同时，推理方向的反对证据也决定了研究边界",
        "在强化学习相关工作负载侧，支持依据来自阶段的天然异质性",
        "但RL工作流的可拆分性不消除一致性和弹性成本",
        "据此，现有研究的共同缺口不是缺少某一种单点调度算法",
    ):
        matches = [p for p in doc.paragraphs if p.text.startswith(old)]
        if matches:
            remove_para(matches[0])

    # Make Chapter 3 explicitly inherit the evidence boundaries from 2.3.
    plan_infer = next(p for p in doc.paragraphs if p.text.startswith("首先通过微基准测试建立设备与链路性能档案"))
    set_text(
        plan_infer,
        "首先通过微基准测试建立设备与链路性能档案，并以外部研究暴露的关键变量作为实验维度：不同batch和上下文长度下的Prefill/Decode吞吐、显存与KV容量、P2P带宽和时延、链路波动、启动与迁移开销。将异构资源池抽象为带属性的图G=(V,E)，把请求、Prefill、Decode、KV传输和同步关系表示为任务图；模型输出推荐协同粒度及候选资源集合，而不是只给出峰值算力排序。",
    )
    plan_infer2 = next(p for p in doc.paragraphs if p.text.startswith("对于长上下文且通信预算允许的请求，再进入HCP协同模块"))
    set_text(
        plan_infer2,
        "对于长上下文且通信预算允许的请求，再进入HCP协同模块。模块根据设备容量和实测计算/通信比生成非均匀序列分片、候选环序及KV块大小，并用代价模型筛除通信主导、尾延迟超标或小设备拖尾严重的配置。实验除TTFT、TPOT和吞吐外，必须报告网络字节数、KV迁移时间、批处理重组、调度决策时间、失败回退和正确性；若任务内协同未超过请求级动态调度，或不满足CloudMatrix式高带宽条件下的通信约束，则记录为边界而非继续扩大部署。",
    )
    plan_rl = next(p for p in doc.paragraphs if p.text.startswith("将LLM-RL后训练表示为包含rollout"))
    set_text(
        plan_rl,
        "将LLM-RL后训练表示为包含rollout、奖励/评估、策略更新和检查点同步的有向任务图，并为每个节点记录模型副本、算力/显存需求、输入输出规模、参数版本和可容忍等待时间。资源分配约束显式区分RLBoost所指出的可抢占、近似无状态rollout与需要紧耦合通信的策略训练；同时纳入队列积压、样本陈旧、参数版本和恢复窗口，使阶段级异构承接不被误写成跨弱链路的常规分布式训练。",
    )
    plan_rl2 = next(p for p in doc.paragraphs if p.text.startswith("在弹性扩展阶段，模拟或实测节点加入退出"))
    set_text(
        plan_rl2,
        "在弹性扩展阶段，模拟或实测节点加入退出、链路带宽下降和检查点恢复等事件。参考DiLoCo、INTELLECT-1和RolloutPipe的边界条件，分别测量重构时间、失效期间有效吞吐、恢复后的参数版本一致性、样本有效性和策略质量变化。去中心化或P2P机制只作为对照路线：若其在资源发现、局部故障恢复或跨域协作上未显著优于集中式/分层基线，便不纳入主路径；多智能体强化学习同样需要先证明局部观测确实造成基线决策瓶颈。",
    )

    reference_template = next(p for p in doc.paragraphs if p.text.startswith("[30]"))
    reference_template._p.addnext(clone_text_paragraph(reference_template, "[32] Xiao Y, Zhou Z, Mao F, et al. An Adaptive Placement and Parallelism Framework for Accelerating RLHF Training[EB/OL]. arXiv:2312.11819, 2023."))
    reference_template._p.addnext(clone_text_paragraph(reference_template, "[31] Bambhaniya A R, Wu H, Subramanian S, et al. MIST: A Co-Design Framework for Heterogeneous, Multi-Stage LLM Inference[EB/OL]. arXiv:2504.09775, 2025."))
    ref32 = next(p for p in doc.paragraphs if p.text.startswith("[32]"))
    new_refs = [
        "[33] He J, Zhai J. FastDecode: High-Throughput GPU-Efficient LLM Serving using Heterogeneous Pipelines[EB/OL]. arXiv:2403.11421, 2024.",
        "[34] Chen L, et al. LLM-PQ: Serving LLM on Heterogeneous Clusters with Phase-Aware Partition and Adaptive Quantization[EB/OL]. arXiv:2403.01136, 2024.",
        "[35] Li X, et al. GoodServe: Towards High-Goodput Serving of Agentic LLM Inferences over Heterogeneous Resources[EB/OL]. arXiv:2605.16867, 2026.",
        "[36] Wei Z, et al. HBM Is Not All You Need: Efficient Disaggregated LLM Serving across Memory-heterogeneous Accelerators[EB/OL]. arXiv:2606.29986, 2026.",
        "[37] Wu Y, Liu X, Zheng H, et al. RLBoost: Harvesting Preemptible Resources for Cost-Efficient Reinforcement Learning on LLMs[EB/OL]. arXiv:2510.19225, 2025.",
        "[38] Chen R, Hu J, Ye K, Xu M. RolloutPipe: Overlapping Pipelined Rollout and Training in Disaggregated On-Policy LLM Reinforcement Learning[EB/OL]. arXiv:2606.26997, 2026.",
    ]
    for ref in reversed(new_refs):
        ref32._p.addnext(clone_text_paragraph(ref32, ref))

    update_toc(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
