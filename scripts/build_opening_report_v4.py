from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path(
    "/Users/stark_sim/Desktop/硕士课题/开题报告/"
    "开题报告_新版2_沈达_面向大模型训练与推理的异构算力协同可行性理论与关键技术研究.docx"
)
OUT = SRC.with_name(
    "开题报告_新版4_沈达_面向大模型推理与强化学习后训练的异构算力协同可行性理论与关键技术研究.docx"
)
NEW_TITLE = "面向大模型推理与强化学习后训练的异构算力协同可行性理论与关键技术研究"


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = deepcopy(para.runs[0]._r.rPr) if para.runs and para.runs[0]._r.rPr is not None else None
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def replace_all(doc, old, new):
    for para in doc.paragraphs:
        if old in para.text:
            set_text(para, para.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        set_text(para, para.text.replace(old, new))


def clone_text_paragraph(template, text):
    p = deepcopy(template._p)
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    r = OxmlElement("w:r")
    if template.runs and template.runs[0]._r.rPr is not None:
        r.append(deepcopy(template.runs[0]._r.rPr))
    t = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def remove_para(para):
    para._element.getparent().remove(para._element)


def insert_after(anchor, heading_template, body_template, heading, body):
    anchor._p.addnext(clone_text_paragraph(body_template, body))
    anchor._p.addnext(clone_text_paragraph(heading_template, heading))


def replace_section(doc, overview, old_start, old_end, heading_template, body_template, items):
    set_text(overview, overview.text)
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if old_start in p.text)
    end = next(i for i, p in enumerate(paragraphs) if old_end in p.text)
    anchor = paragraphs[start - 1]
    # Each replaced heading has one body paragraph. Remove the final body too.
    for p in paragraphs[start : end + 2][::-1]:
        remove_para(p)
    for heading, body in reversed(items):
        anchor._p.addnext(clone_text_paragraph(body_template, body))
        anchor._p.addnext(clone_text_paragraph(heading_template, heading))


def set_toc_label(para, label):
    texts = list(para.iter(qn("w:t")))
    if not texts:
        raise ValueError("TOC paragraph has no visible text")
    texts[0].text = label


def update_toc(doc):
    toc = doc.part._element.xpath(".//w:sdt//w:p")
    if len(toc) < 24:
        raise ValueError("table of contents not found")

    set_toc_label(toc[15], "3.1.1  异构推理的多层次协同扩展")
    set_toc_label(toc[16], "3.1.2  强化学习工作负载的异构承接")
    set_toc_label(toc[20], "3.2.1  异构推理协同扩展方案")
    set_toc_label(toc[21], "3.2.2  强化学习负载弹性编排方案")

    for index in (23, 22, 18, 17):
        para = toc[index]
        para.getparent().remove(para)

    update_fields = doc.settings.element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        doc.settings.element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main():
    doc = Document(SRC)

    replace_all(doc, "面向大模型训练与推理的异构算力协同可行性理论与关键技术研究", NEW_TITLE)
    replace_all(doc, "大模型训练与推理", "大模型推理与强化学习后训练")
    replace_all(doc, "大模型训练和推理", "大模型推理与强化学习后训练")
    replace_all(doc, "大模型训练与推理任务", "大模型推理与强化学习后训练任务")
    replace_all(doc, "典型大模型训练与推理任务", "典型大模型推理与强化学习后训练任务")
    replace_all(doc, "不同大模型训练与推理任务", "不同大模型推理与强化学习后训练任务")
    replace_all(doc, "分布式大模型训练与推理系统", "大模型推理与 LLM-RL 后训练系统")
    replace_all(doc, "大模型训练与推理系统", "大模型推理与 LLM-RL 后训练系统")
    replace_all(doc, "分布式训练和其他任务", "LLM-RL 后训练工作流和其他阶段性任务")
    replace_all(doc, "分布式训练、强化学习工作流或去中心化算力协作", "LLM-RL 后训练工作流或去中心化弹性协作")
    replace_all(doc, "异构分布式训练与推理系统调研", "异构推理与 LLM-RL 后训练系统调研")
    replace_all(doc, "异构训练、异构推理", "异构推理与 LLM-RL 后训练")
    replace_all(doc, "分布式训练和去中心化强化学习方向", "LLM-RL 后训练与去中心化协同方向")
    replace_all(doc, "训练、监督微调、强化学习后训练和去中心化训练", "推理与 LLM-RL 后训练")
    replace_all(doc, "大模型训练和推理", "大模型推理与 LLM-RL 后训练")
    replace_all(doc, "大模型服务与后训练", "大模型推理服务与 LLM-RL 后训练")
    replace_all(doc, "大模型训练任务", "LLM-RL 后训练任务")
    replace_all(doc, "异构训练系统", "异构大模型系统")

    current_route_para = next(p for p in doc.paragraphs if "目前针对异构环境的主流思路" in p.text)
    set_text(
        current_route_para,
        "目前针对异构环境的主流思路，是将具有不同资源需求的请求、阶段或工作流任务分配给相匹配的设备。推理系统中出现了 Prefill 与 Decode 分离部署、面向异构资源池的请求路由等设计；强化学习工作流中则可将样本生成、奖励评估、策略更新等阶段分配给性能匹配的设备组。这些方法说明异构资源可以在服务和工作流层形成互补，但还缺少一套依据负载可拆分性、通信强度、设备能力和服务目标选择协同层次的统一方法。特别是张量并行、数据并行等高频同步训练方式对跨设备通信要求极高，不能简单假设异构设备共同承担同一计算即可获得收益；而长上下文 Prefill 等特定场景则可能在通信预算可接受时采用更细粒度的协同。",
    )

    source_para = next(p for p in doc.paragraphs if "本课题即源于" in p.text)
    set_text(
        source_para,
        "本课题即源于对上述问题的思考与前期实践。申请人围绕异构环境下基于环结构的上下文并行通信开展研究，已完成 hetero-cp-ringattn 原型系统的核心代码开发，并在小型混合集群上获得了初步加速效果。在此基础上，课题以大模型推理与 LLM-RL 后训练系统为主要对象：在推理侧参考 Dynamo 的请求编排、Prefill/Decode 分离和可插拔执行接口，探索异构资源池与 HCP 类细粒度协同模块的衔接；在强化学习侧参考 HetRL 的异构工作流调度，以及 Prime RL 和 Prime DiLoCo 在异步后训练、弹性设备网格和弱同步通信方面的实践，研究可拆分负载在异构资源上的承接与编排。",
    )

    layout_para = next(p for p in doc.paragraphs if "从研究布局来看" in p.text)
    set_text(
        layout_para,
        "从研究布局来看，现有工作呈现明显的层次割裂。Dynamo 等工业系统主要关注大模型推理服务的资源编排、Prefill/Decode 分离、KV 传输和执行引擎组织，其价值在于提供请求级和阶段级调度抽象，而非已经解决异构设备共同执行单一任务的问题；HetRL 关注包含 rollout、奖励评估、策略更新等阶段的异构 LLM-RL 后训练工作流；Prime RL 和 Prime DiLoCo 则展示了异步后训练、弹性设备网格、弱同步通信、节点动态加入退出和跨地域协作的工程可行性。这些工作分别覆盖服务层、工作流层和通信基础设施层，但尚缺少依据通信约束选择协同粒度，并把异构能力从资源池稳健地扩展到推理执行和强化学习工作流的方法。",
    )

    scale_para = next(p for p in doc.paragraphs if "规模层面的矛盾同样突出" in p.text)
    set_text(
        scale_para,
        "规模层面的矛盾同样突出。现有异构系统普遍沿用集中式调度架构，隐含全局拓扑可见、链路均匀可靠的假设；当节点规模扩大、链路质量参差不齐时，单一调度器的状态采集与决策开销会增加。Prime DiLoCo 的 ElasticDeviceMesh、心跳失效检测、动态进程组调整、异步检查点和弱同步通信，说明在 LLM-RL 后训练中可以把节点弹性和弱连接协作作为工程扩展；Prime RL 则将异步 rollout、训练、评估、FSDP/EP/CP 和 vLLM 执行引擎整合到统一工作流中。因此，去中心化并非本课题的既定前提，而是在跨域资源发现、局部故障恢复或弱连接协作成为瓶颈时可选的支撑机制；核心仍是针对设备能力和通信条件建立可解释的协同决策方法。",
    )

    value_para = next(p for p in doc.paragraphs if "就研究价值而言" in p.text)
    set_text(
        value_para,
        "就研究价值而言，本课题首先希望建立适用于大模型推理和 LLM-RL 后训练的异构算力协同可行性判定模型，为给定的设备组合、网络条件和任务图选择请求级、阶段级或任务内协同的合适粒度，并给出资源分配与效率边界。其次，课题将在技术层面打通从建模到系统的通路：在推理场景中探索把 HCP 的容量感知非均匀协同作为可插拔模块接入 vLLM 或 Dynamo 类执行框架，在通信预算可接受的长上下文 Prefill 中进行细粒度验证；在 LLM-RL 后训练场景中研究 rollout、评估和策略更新等可拆分阶段的异构承接、异步流水线和弹性编排。",
    )

    review_para = next(p for p in doc.paragraphs if "从上述进展可以看出" in p.text)
    set_text(
        review_para,
        "从上述进展可以看出，异构协同的关键不在于强行让所有设备共同执行高通信耦合的单一任务，而在于依据负载可拆分性和通信约束选择协同粒度：资源池和请求级调度适合大多数在线推理负载，阶段级编排适合强化学习后训练工作流，只有在长上下文 Prefill 等通信预算可接受的场景中，才进一步探索 HCP 类任务内协同。去中心化拓扑可为跨域资源发现、容错和弱连接协作提供补充，但不作为每类负载都必须采用的机制。",
    )

    direction_para = next(p for p in doc.paragraphs if "基于以上分析，本课题拟以" in p.text)
    set_text(
        direction_para,
        "基于以上分析，本课题拟以异构负载的分层协同扩展为主线，从可行性建模入手，分别在推理服务和强化学习工作流中验证请求级、阶段级与条件性任务内协同的适用边界，并在确有需要时引入弱中心或去中心化支撑机制，形成面向异构大模型基础设施的可验证技术方案。",
    )

    # 3.1: five small points -> two load-oriented research points.
    overview = next(p for p in doc.paragraphs if "本课题以异构算力协同的可行性" in p.text)
    set_text(
        overview,
        "本课题围绕异构算力协同的可行性与可持续性，研究对象限定为两类具有代表性的负载：一是大模型在线推理，研究从异构资源池、请求路由到条件性任务内协同的扩展路径；二是强化学习相关工作负载，以 LLM-RL 后训练为主要实例，研究 rollout、奖励评估和策略更新等可拆分阶段的异构承接。两点共享设备—网络—任务可行性建模和容量感知调度方法，但不预设所有异构设备均需共同执行一个高通信耦合任务。去中心化和多智能体强化学习仅在集中式或分层基线无法满足跨域资源发现、容错或动态决策需求时作为可选扩展。",
    )
    h_template = next(p for p in doc.paragraphs if "3.1.1" in p.text)
    b_template = next(p for p in doc.paragraphs if "研究面向异构算力协同的统一性能" in p.text)
    items = [
        (
            "3.1.1  面向异构推理的多层次协同扩展",
            "面向大模型在线推理研究异构能力从资源池到执行层的渐进式扩展。首先刻画设备计算速率、显存容量、链路带宽、启动开销和可靠性，建立请求级、阶段级与任务内协同的可行性和代价模型；其次参考 Dynamo 的 Prefill/Decode 分离、请求编排、KV 传输和可插拔执行接口，研究异构设备上的请求路由、阶段放置、资源配比和计算—通信联合调度；最后以长上下文 Prefill 为条件性验证场景，探索将 hetero-cp-ringattn 的容量感知非均匀切分、环序和 KV 块传输封装为可接入 vLLM 或 Dynamo 类框架的协同模块。HCP 只在通信预算与正确性约束满足时用于任务内协同，不将异构设备共同完成单一推理任务作为普遍前提。",
        ),
        (
            "3.1.2  面向强化学习工作负载的异构承接与弹性编排",
            "面向强化学习相关工作负载研究可拆分阶段在异构资源上的承接、编排与弹性执行，以 LLM-RL 后训练为主要验证实例。围绕 rollout、奖励计算、评估、策略更新和检查点等阶段，研究设备能力感知的任务放置、依赖编排、并发度调节和结果/参数传递；参考 HetRL 的阶段级异构调度，结合 Prime RL 的异步执行以及 Prime DiLoCo 的弹性设备网格、弱同步通信、节点动态加入退出和异步检查点，研究弱连接环境下的吞吐、容错与资源利用率。去中心化仅作为跨域资源发现、局部故障恢复或弱中心协作的可选扩展；多智能体强化学习也仅作为在线调度候选方法，需在启发式与约束优化基线不足时再引入，且不与被调度的 LLM-RL 负载混同。",
        ),
    ]
    replace_section(doc, overview, "3.1.1", "3.1.5", h_template, b_template, items)

    # 3.2: five implementation schemes -> two corresponding schemes.
    scheme = next(p for p in doc.paragraphs if p.text.strip() == "3.2  研究方案")
    h_template = next(p for p in doc.paragraphs if "3.2.1" in p.text)
    b_template = next(p for p in doc.paragraphs if "首先通过微基准测试建立设备" in p.text)
    items = [
        (
            "3.2.1  异构推理协同扩展方案",
            "首先通过微基准测试建立设备与链路性能档案，将异构资源池抽象为带属性的图 G=(V,E)，并把请求、Prefill、Decode、KV 传输和同步关系表示为任务图。其次以集中式或分层调度为基线：上层完成候选设备、SLO 与成本感知的请求路由，中层决定 Prefill/Decode 阶段放置、资源配比和流水化参数，底层通过可插拔后端执行模型服务与 KV 传输。对于长上下文且通信预算允许的请求，再调用 HCP 协同模块决定非均匀序列分片、环序和 KV 块传输；验证其正确性、计算通信重叠和 P2P 回退路径。对于跨域节点失效或链路抖动，优先采用局部状态上报与回退路由，必要时再研究局部邻居视图和逻辑环重构。实验比较单设备、静态异构资源池、动态请求调度与条件性 HCP 扩展，报告 TTFT、吞吐、显存峰值、通信开销和恢复时间。",
        ),
        (
            "3.2.2  强化学习工作负载的异构承接与弹性编排方案",
            "将 LLM-RL 后训练表示为包含 rollout、奖励/评估、策略更新和检查点同步的有向任务图，结合设备画像、网络状态和阶段依赖建立资源分配与可行性约束。参考 HetRL 的异构阶段级调度方式，研究不同设备组在各阶段的放置、并发和结果/参数传递；参考 Prime RL 的异步 rollout—训练—评估流水线，研究在集中式或分层协调下的稳定基线。对于跨域节点加入退出、局部失效和弱连接等情形，再参考 Prime DiLoCo 的 ElasticDeviceMesh、心跳检测、弱同步通信和异步检查点，评估弱中心或 P2P 扩展的收益与代价。动态决策先采用启发式与约束优化，只有当局部观测和状态变化使其无法有效决策时，才引入多智能体强化学习并保留安全回退策略。实验重点比较阶段级异构编排、静态同构基线、动态调度与条件性弹性协同在吞吐、策略更新延迟、资源利用率、通信量和恢复时间上的差异。",
        ),
    ]
    replace_section(doc, scheme, "3.2.1", "3.2.5", h_template, b_template, items)

    # Objectives and schedule should expose the same two-load boundary.
    for p in doc.paragraphs:
        if "建立设备—网络—任务统一的异构算力协同模型" in p.text:
            set_text(p, "（1）理论目标：建立设备—网络—任务统一的异构算力协同模型，给出面向大模型推理与 LLM-RL 后训练的协同粒度选择、可行性判定、资源约束表达和效率边界分析。")
        elif "形成容量感知的非均匀任务划分" in p.text:
            set_text(p, "（2）方法目标：形成面向推理与 LLM-RL 后训练的容量感知路由、阶段编排、条件性任务内协同、计算—通信联合调度和局部容错回退方法，并明确集中式、弱中心和多智能体决策的适用边界。")
        elif "并为分布式训练、强化学习工作流和其他任务保留统一接口" in p.text:
            set_text(p, "（3）系统目标：实现可插拔的异构协同原型系统，以 hetero-cp-ringattn 为长上下文推理的条件性协同载体，探索其与 vLLM 或 Dynamo 类框架的衔接，并为 LLM-RL 后训练工作流保留阶段级异构承接接口。")
        elif "根据前期结果选择" in p.text and "扩展验证场景" in p.text:
            set_text(p, "（4）实验目标：在多代、多厂商或不同设备能力的混合环境中，对比单设备、同构基线、静态异构编排、动态调度和条件性协同方法，系统评估推理与 LLM-RL 后训练的完成时间、吞吐、显存、通信、扩展性和故障恢复能力。")
        elif "不预设单一固定加速倍数" in p.text:
            set_text(p, "（4）实验目标：在多代、多厂商或不同设备能力的混合环境中，对比单设备、同构基线、静态异构编排、动态调度和条件性协同方法，系统评估推理与 LLM-RL 后训练的完成时间、吞吐、显存、通信、扩展性和故障恢复能力；不预设单一固定加速倍数，而是报告不同条件下的收益区间和失效边界。")
        elif "围绕统一可行性建模" in p.text:
            set_text(p, "（5）成果目标：围绕协同粒度可行性建模、异构推理扩展、强化学习工作流编排和条件性弹性机制形成可拆分的论文与专利成果，持续维护并开源可复用的系统代码和实验基准。")
        elif "已完成对异构推理与" in p.text:
            set_text(p, "已完成对异构推理与 LLM-RL 后训练、上下文并行、Prefill/Decode 分离及弹性协同系统的调研，重点分析异构请求路由、阶段级承接、P2P 协同和多智能体动态调度的适用条件，初步明确请求级、阶段级与条件性任务内协同的研究空间。")
        elif "异构设备画像与去中心化协同探索" in p.text:
            set_text(p, "5.1.1.3 异构设备画像与弹性协同探索")
        elif "已围绕设备算力、显存容量、链路带宽和通信延迟" in p.text:
            set_text(p, "已围绕设备算力、显存容量、链路带宽和通信延迟开展初步测量，完成动态环构建、非均匀切片和局部拓扑变化的仿真探索，并保留跨 CUDA、ROCm 与 MPS 后端协同验证所需的接口设计。后续将把这些工程结果统一纳入协同粒度可行性模型，并以集中式或分层调度为基线评估弹性扩展的必要性。")
        elif "形成 hetero-cp-ringattn 开源代码库" in p.text:
            set_text(p, "（1）形成 hetero-cp-ringattn 开源代码库和异构通信原型；（2）完成异构推理、LLM-RL 后训练、上下文并行与弹性协同方向的文献和系统调研；（3）搭建异构协同仿真与实验环境，具备开展设备画像、通信性能、条件性任务内协同和端到端推理验证的条件；（4）围绕异构请求调度、通信机制和工作流承接形成阶段性技术报告，并将根据后续实验结果拆分论文和专利选题。")
        elif "去中心化环境下通信环的频繁重构" in p.text:
            set_text(p, "在跨域或弱连接条件下，局部重构和 P2P 扩展可能引发路由抖动，进而影响请求或工作流的时延稳定性。对此，课题将默认采用集中式或分层调度与回退路由；只有基线在资源发现、故障恢复或弱连接协作上不足时，才采用局部一致性协议、版本化握手和短暂双轨过渡等机制，并单独评估其收益与开销。")
        elif "去中心化多智能体强化学习面对的状态空间" in p.text:
            set_text(p, "多智能体强化学习若用于异构调度，可能面对状态空间庞大、奖励信号稀疏和训练初期收敛缓慢的问题。因此本课题不将其作为既定实现路线，而是先以启发式和约束优化建立可解释的稳定基线；仅当局部观测和动态变化使基线无法满足需求时，再在小规模场景采用课程学习和集中式训练、分散式执行进行对照验证，并保留回退策略。")

    for row in doc.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "根据前期结果选择" in p.text:
                    set_text(p, "根据前期结果完成 LLM-RL 后训练工作流的扩展验证；仅在跨域资源发现或弱连接协作确有收益时，补充去中心化弹性协作实验。")
                elif "研究多智能体强化学习动态适应" in p.text:
                    set_text(p, "完成异构推理的端到端验证；比较静态编排、动态调度与条件性 HCP 扩展。若基线暴露出局部观测决策瓶颈，再开展多智能体强化学习对照实验。")
                elif "研究容量感知任务划分、计算" in p.text:
                    set_text(p, "研究容量感知请求路由、阶段放置、计算—通信联合调度和 P2P 回退机制；完成 vLLM 或 Dynamo 类框架衔接的原型设计。")

    update_toc(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
