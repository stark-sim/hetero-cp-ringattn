from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path('/Users/stark_sim/Desktop/硕士课题/开题报告/开题报告_新版9_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx')
OUT = SRC.with_name('开题报告_新版10_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx')


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = deepcopy(para.runs[0]._r.rPr) if para.runs and para.runs[0]._r.rPr is not None else None
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def remove_para(para):
    para._element.getparent().remove(para._element)


def first(doc, pred):
    return next(p for p in doc.paragraphs if pred(p.text))


def insert_after(para, text):
    new_p = deepcopy(para._p)
    for child in list(new_p):
        if child.tag != qn('w:pPr'):
            new_p.remove(child)
    para._p.addnext(new_p)
    new_para = next(p for p in para._parent.paragraphs if p._p is new_p)
    set_text(new_para, text)
    return new_para


def insert_after_with_template(para, text, template):
    new_p = deepcopy(template._p)
    para._p.addnext(new_p)
    new_para = next(p for p in para._parent.paragraphs if p._p is new_p)
    set_text(new_para, text)
    return new_para


def replace_all(doc, old, new):
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))


def update_toc(doc):
    replacements = {
        '2.1.1  多任务异构调度': '2.1.1  异构推理服务的请求与阶段承接',
        '2.1.2  单任务异构协同': '2.1.2  长上下文Prefill的任务内异构协同',
        '2.1.3  去中心化环境下的异构协同': '2.1.3  LLM-RL后训练的异构阶段承接与弹性执行',
        '2.1.4  去中心化与弱连接协作的适用边界': '2.1.4  去中心化与弱连接协作的适用边界',
        '2.2.1  异构计算负载均衡与调度': '2.2.1  异构推理系统与资源编排',
        '2.2.2  异构推理系统优化': '2.2.2  LLM-RL工作流与异构资源支持',
        '3. 主要研究内容及研究方案': '3. 主要研究内容及实施方案',
        '3.1  主要研究内容': '3.1  研究内容',
        '3.2  研究方案': '3.2  实施方案',
    }
    for p in doc.part._element.xpath('.//w:sdt//w:p'):
        text = ''.join(t.text or '' for t in p.iter(qn('w:t')))
        for old, new in replacements.items():
            if old in text:
                ts = list(p.iter(qn('w:t')))
                # TOC entries commonly split the heading across two runs and
                # keep the page number in the final run. Replace only heading
                # runs so page fields and their formatting remain untouched.
                page_run = ts[-1] if ts and (ts[-1].text or '').strip().isdigit() else None
                heading_ts = ts[:-1] if page_run is not None else ts
                if heading_ts:
                    heading_ts[0].text = new
                    for t in heading_ts[1:]:
                        t.text = ''
                break
    # The source TOC has no 2.1.4 row. Clone the existing 2.1.3 row so the
    # new boundary subsection appears with the same tabs and page-number run.
    toc_ps = doc.part._element.xpath('.//w:sdt//w:p')
    marker = None
    for p in toc_ps:
        text = ''.join(t.text or '' for t in p.iter(qn('w:t')))
        if text.startswith('2.1.3  LLM-RL后训练的异构阶段承接与弹性执行'):
            marker = p
            break
    if marker is not None and not any(
        ''.join(t.text or '' for t in p.iter(qn('w:t'))).startswith('2.1.4  去中心化与弱连接协作的适用边界')
        for p in toc_ps
    ):
        clone = deepcopy(marker)
        ts = list(clone.iter(qn('w:t')))
        page_run = ts[-1] if ts and (ts[-1].text or '').strip().isdigit() else None
        heading_ts = ts[:-1] if page_run is not None else ts
        if heading_ts:
            heading_ts[0].text = '2.1.4  去中心化与弱连接协作的适用边界'
            for t in heading_ts[1:]:
                t.text = ''
        if page_run is not None:
            page_run.text = '4'
        marker.addnext(clone)


def main():
    doc = Document(SRC)

    # Chapter 1: define the actual workload scope and keep broad claims out.
    set_text(first(doc, lambda t: t.startswith('近年来，大语言模型与强化学习的规模化发展')),
        '近年来，大语言模型推理和LLM-RL后训练对算力的需求持续增长：模型上下文变长，在线请求具有明显的Prefill与Decode阶段差异，后训练工作流又包含rollout、奖励计算、评估和策略更新等不同阶段。与此同时，数据中心和实验环境中长期存在不同厂商、不同代际、不同显存容量的加速卡。高端设备紧张与存量设备闲置并存，使“哪些工作负载能够由异构卡组合承接、在什么条件下值得承接”成为一个比简单追求多卡并行更具体的系统问题。')
    set_text(first(doc, lambda t: t.startswith('目前针对异构环境的主流思路')),
        '现有异构系统通常先利用负载之间的可分性：推理服务通过请求路由或Prefill/Decode分离把不同阶段放到匹配的设备，LLM-RL后训练则把rollout、奖励评估和数据处理等阶段分配给不同设备组。对于长上下文Prefill，若单设备显存不足且链路能够承受K/V传递，还可以进一步研究一次注意力计算内部的非对称上下文并行。相反，张量并行、数据并行和策略更新往往需要高频同步，不能因为设备“能够连接”就假定跨弱链路组合会带来收益。')
    set_text(first(doc, lambda t: t.startswith('本课题即源于对上述问题的思考')),
        '本课题即源于对上述问题的思考与前期实践。申请人已完成hetero-cp-ringattn原型的核心开发，围绕非均匀上下文切片、K/V P2P ring和online softmax开展了跨设备验证。后续研究固定为两条互相独立但共享画像方法的路线：推理侧以长上下文Prefill为对象，在vLLM或Dynamo类服务接口下研究HCP准入、非对称切分和回退；LLM-RL侧以rollout、奖励计算、评估和数据处理等阶段为对象，参考HetRL、Prime RL和Prime DiLoCo研究异构放置、异步流水线与弹性恢复。课题不把通用大模型预训练、跨弱链路TP/DP或HCP参与策略更新列为目标。')

    set_text(first(doc, lambda t: t.startswith('与同构集群相比，异构集群的根本特征')),
        '与同构集群相比，异构集群在设备算力、显存容量、后端兼容性以及设备间链路带宽和时延方面存在差异。差异本身并不意味着所有任务都适合合作：如果任务需要频繁交换激活或梯度，最慢设备和最弱链路可能决定整体速度；如果任务能够按请求或阶段解耦，异构设备则可以通过分工提高资源利用率。本课题因此把“负载可拆分性—通信强度—设备能力—服务目标”作为判断异构承接的基本约束。')
    set_text(first(doc, lambda t: t.startswith('从研究布局来看，现有工作呈现明显的层次割裂')),
        '从研究布局来看，现有工作已经分别覆盖了若干可复用的层次。Dynamo、vLLM及相关系统提供请求编排、Prefill/Decode分离、KV管理和执行接口；HetRL关注LLM-RL阶段在异构GPU上的放置；Prime RL提供异步rollout—训练—评估工作流，Prime DiLoCo提供弹性设备网格、弱同步和检查点机制。这些工作并未共同解决“异构卡如何承接长上下文Prefill中的一次任务内计算”或“阶段承接如何在版本、队列和故障约束下保持有效”这两个具体问题，因而为本课题留下了明确的增量空间。')
    set_text(first(doc, lambda t: t.startswith('规模层面的矛盾同样突出')),
        '在资源动态变化时，异构承接还要面对准入、排队、版本和恢复问题。推理请求若在执行中遇到链路抖动或小设备拖尾，需要能够回退到单设备或请求/阶段级路径；LLM-RL工作流若出现节点退出，则必须判断哪些rollout或评估任务可以迁移，哪些策略更新必须等待一致版本。去中心化发现和多智能体决策可以作为局部机制进行对照，但不是本课题预先承诺的系统形态。')
    set_text(first(doc, lambda t: t.startswith('就研究价值而言，本课题首先希望建立')),
        '本课题的研究价值不在于提出一个覆盖所有AI任务的异构大框架，而在于形成可执行的负载承接边界。对推理侧，研究HCP在长上下文Prefill中的容量感知非对称协同，并明确它与vLLM/Dynamo已有请求和阶段调度的分工；对LLM-RL侧，研究rollout、奖励评估和数据处理等阶段的异构放置、版本/队列控制与弹性恢复，并明确策略更新保留在高带宽同构或局部同构设备组的条件。')
    set_text(first(doc, lambda t: t.startswith('这些工作的实际意义在于直接服务')),
        '这些工作的实际意义在于让异构资源以可预测的方式承接真实工作，而不是以“参与设备越多越好”为成功标准。若中低端或不同厂商设备能够在给定上下文长度、网络条件和RL阶段比例下贡献有效吞吐，系统可减少对单一高端设备的依赖；若通信、拖尾或恢复代价超过收益，准入规则也应明确拒绝该组合，为异构资源池的实际部署提供负向边界。')

    # Chapter 2: reorganize the evidence around the two concrete research lines.
    set_text(first(doc, lambda t: t.strip() == '2.1.1  多任务异构调度'), '2.1.1  异构推理服务的请求与阶段承接')
    set_text(first(doc, lambda t: t.startswith('围绕异构GPU的利用问题')),
        '国外推理系统首先在请求和阶段层利用异构性。DistServe、Splitwise和Mooncake将Prefill与Decode分离并分别扩展资源池；GoodServe、ThunderServe和MIST进一步把设备性能、成本、抢占和服务目标纳入放置或重调度；Dynamo、vLLM生态及llm-d则把请求路由、KV生命周期和执行后端组织成可插拔的服务接口。它们共同说明，异构卡可以通过请求级或阶段级分工承接在线推理，但这些框架本身并不等于异构设备共同完成同一份注意力计算。')
    set_text(first(doc, lambda t: t.startswith('在推理服务方向，DistServe')),
        '这类工作的直接启示是：异构推理的基线应先包含单设备、请求路由和Prefill/Decode分离，再判断是否需要更细粒度协同。其不足也很明确：请求级放置无法解决单个超长请求的显存墙，阶段级分离还会引入KV传输、排队重组和尾延迟；因此，HCP只有在上下文长度、显存容量和链路条件同时满足时才有研究价值。')
    set_text(first(doc, lambda t: t.startswith('训练方向的工作同样以任务划分为核心')),
        '在相关训练与服务研究中，HexiScale、HeterMoE和HexGen-2分别从非对称并行、专家/注意力分工和异构副本组织角度扩大了可承接范围，但它们关注的训练并行或服务图划分与本课题的长上下文Prefill并不相同。故本课题将这些工作作为“异构性应进入放置与切分决策”的证据，而不把其训练结果直接外推到推理。')
    set_text(first(doc, lambda t: t.strip() == '2.1.2  单任务异构协同'), '2.1.2  长上下文Prefill的任务内异构协同')
    set_text(first(doc, lambda t: t.startswith('相比之下，聚焦于单个任务内部')),
        '任务内异构协同的研究数量相对较少，核心问题是如何在不均匀设备和非均匀链路下保持一次计算的正确性与可接受效率。Ring Attention以环形传递K/V块避免一次性收集完整上下文，为长序列注意力的分块执行提供了基础；USP讨论了统一序列并行与混合通信模式；HexiSeq进一步针对异构训练提出非均匀序列分片和加权分配。')
    set_text(first(doc, lambda t: t.startswith('HexiSeq')),
        'HexiSeq的非均匀分片与加权分配直接说明均分假设可能使小显存或低算力设备先成为瓶颈，但其对象主要是训练。Hetis在推理侧采用角色分化的辅助协作，更多是机会性加速；USP虽讨论混合通信，却未把设备能力差异显式用于Prefill负载切分。由此可见，长上下文推理中的非对称HCP仍需要单独验证。')
    set_text(first(doc, lambda t: t.startswith('Hetis')),
        '现有任务内方案的共同限制是对目标阶段、通信协议或硬件假设较为特定：训练方案不能直接证明推理Prefill的端到端收益，解码辅助方案也不能替代一次Prefill中的完整K/V协同。实际系统还必须处理显存瞬态工作区、P2P带宽波动、ring拖尾和失败回退，这正是本课题拟在HCP原型上补充的内容。')
    set_text(first(doc, lambda t: t.startswith('USP')),
        '因此，本课题不把“异构设备能够共同执行一次注意力”作为普遍结论，而把它设为有准入条件的研究假设：只有当非对称seq_chunk_len能够解决显存承接问题，且K/V传输和在线合并的代价不超过单设备或请求/阶段路由的收益时，才认为任务内协同成立。')

    # Replace the former decentralization section with the LLM-RL evidence line.
    set_text(first(doc, lambda t: t.strip() == '2.1.3  去中心化环境下的异构协同'), '2.1.3  LLM-RL后训练的异构阶段承接与弹性执行')
    set_text(first(doc, lambda t: t.startswith('去中心化技术与异构计算的结合')),
        'LLM-RL后训练天然包含多个阶段，异构性更适合首先在工作流层体现。HetRL将rollout、奖励评估和策略更新作为不同资源需求的调度单元；自适应RLHF放置、RLBoost和RolloutPipe分别从设备放置、可抢占rollout资源和流水化执行角度说明，生成与评估类任务可以与策略更新解耦。')
    set_text(first(doc, lambda t: t.startswith('PRIME-RL')),
        'Prime RL的公开工作流把异步rollout、训练、评估以及vLLM等执行引擎组织在同一后训练管线中，证明异构设备可以通过阶段级和异步方式共同提供有效工作。Prime DiLoCo的ElasticDeviceMesh、异步检查点和节点弹性则提供了动态加入退出与弱同步的工程参考。上述系统已经具备相当多的基础能力，本课题不重复实现其通用训练框架，而聚焦阶段能力合同、版本/队列约束、样本完整性和局部回退等组合条件。')
    set_text(first(doc, lambda t: t.startswith('Sailor')),
        '这条证据线也有明确反面约束：策略更新通常需要较高带宽和较强一致性，异步rollout会引入版本滞后和样本陈旧，节点变化会增加检查点与恢复开销。因此，本课题把策略更新默认保留在高带宽同构或局部同构设备组，只把rollout、奖励计算、评估和数据处理作为主要异构承接对象，并以有效样本成本和策略质量而非原始设备利用率作为评价依据。')
    set_text(first(doc, lambda t: t.startswith('PCCL')),
        '综合来看，LLM-RL异构支持的研究缺口不是“再造一个RL算法”，而是把阶段放置、异步执行和弹性恢复放进同一组可验证的能力约束中：设备必须能运行目标后端和模型，队列与版本必须可控，节点失效后必须有可测量的恢复路径。')

    # Add a separate, compact subsection for decentralization instead of making it a main line.
    p = first(doc, lambda t: t.startswith('从上述进展可以看出，异构协同的关键'))
    set_text(p, '2.1.4  去中心化与弱连接协作的适用边界')
    heading_213 = first(doc, lambda t: t.startswith('2.1.3  LLM-RL后训练的异构阶段承接与弹性执行'))
    body_template = first(doc, lambda t: t.startswith('LLM-RL后训练天然包含多个阶段'))
    insert_after_with_template(p, 'DiLoCo、INTELLECT-1、DisTrO、Sailor和Prime DiLoCo等工作表明，低频同步、压缩通信、P2P发现或弹性设备网格能够支持跨地域和动态节点协作；PCCL等工作则进一步关注不稳定公网中的容错通信。这些成果支持把去中心化机制作为资源发现、局部恢复和弱连接协作的候选工具，但它们主要作用于工作流或参数更新层，并没有证明其适合长上下文Prefill中的紧耦合K/V ring。故本课题只在集中式或分层基线出现全局状态采集、跨域发现或局部恢复瓶颈时做小规模对照，不把去中心化和多智能体强化学习列为独立研究点。', body_template)
    # Match the newly introduced subsection heading to the neighboring heading.
    if p._p.pPr is not None:
        p._p.remove(p._p.pPr)
    if heading_213._p.pPr is not None:
        p._p.insert(0, deepcopy(heading_213._p.pPr))
    if p.runs and heading_213.runs and heading_213.runs[0]._r.rPr is not None:
        p.runs[0]._r.insert(0, deepcopy(heading_213.runs[0]._r.rPr))
    # Keep the old paragraph slot after the inserted heading as a domestic section.
    set_text(first(doc, lambda t: t.strip() == '2.2  国内研究现状'), '2.2  国内研究现状')
    set_text(first(doc, lambda t: t.strip() == '2.2.1  异构计算负载均衡与调度'), '2.2.1  异构推理系统与资源编排')
    set_text(first(doc, lambda t: t.startswith('国内的异构大模型系统研究在训练侧')),
        '国内相关工作在异构推理的资源编排和系统优化方面持续推进。HexGen-2通过图划分与约束优化组织异构GPU副本和并行策略；FlowKV围绕KV Cache跨机传输优化内存分配器和传输流水线；针对显存层次差异的分解式部署工作进一步说明，模型权重、KV缓存和请求状态可以采用不同的资源组织方式。它们为本课题提供请求/阶段路由和KV管理的基线，但尚未直接给出HCP在长上下文Prefill中的非对称任务内计划。')
    set_text(first(doc, lambda t: t.strip() == '2.2.2  异构推理系统优化'), '2.2.2  LLM-RL工作流与异构资源支持')
    set_text(first(doc, lambda t: t.startswith('推理侧的代表工作是上海交通大学')),
        '国内关于LLM-RL的公开研究和工程实践更多集中在RLHF放置、流水化和资源利用问题。自适应RLHF放置工作将不同阶段映射到不同设备；RLBoost考察抢占式资源对rollout的承接；RolloutPipe研究rollout与训练的流水化重叠。这些工作支持“生成和评估阶段适合利用异构资源”的判断，但也暴露出版本滞后、样本有效性和恢复成本等需要系统约束的问题。与本课题相关的异构训练框架如HETHUB、AutoHet可作为训练侧对照，不能直接替代LLM-RL阶段承接实验。')

    # 2.3: explicit support, counter-evidence, and the resulting scope.
    p72 = first(doc, lambda t: t.startswith('从证据类型看，异构推理和异构RL'))
    p73 = first(doc, lambda t: t.startswith('对于异构推理，支持证据'))
    p74 = first(doc, lambda t: t.startswith('对于异构RL，支持证据'))
    p75 = first(doc, lambda t: t.startswith('综合来看，现有研究的共同缺口'))
    p76 = first(doc, lambda t: t.startswith('宏观层面的成本与能源约束'))
    set_text(p72, '从现有证据看，异构推理和异构RL后训练具有相关但不相同的承接逻辑。推理侧已有请求路由、阶段分离、KV迁移和部分任务内并行的实验；RL侧的直接证据主要来自阶段放置、异步rollout、弹性设备网格和低频同步。两条证据线共同支持按负载特征分工，但都不支持异构设备无条件共同执行任意高通信计算。')
    set_text(p73, '推理侧的支持观点是：请求输出长度、KV命中、显存容量和设备速度差异会改变最优放置，长上下文Prefill还存在单设备显存墙，因此容量感知的非对称切分可能带来额外承接能力。反对观点是：KV迁移、ring拖尾、批处理重组、链路波动和尾延迟可能抵消计算收益，HCP必须以准入和回退为前提。')
    set_text(p74, 'LLM-RL侧的支持观点是：rollout、奖励评估、评估和数据处理具有阶段边界，异步流水线与弹性资源能够提高有效资源利用率，HetRL、Prime RL、RLBoost和RolloutPipe提供了可复用的工作流依据。反对观点是：策略更新和参数同步仍然对带宽与一致性敏感，异步会造成版本滞后、样本陈旧和恢复开销，因此不能把整个RL训练过程简单地异构混布。')
    set_text(p75, '由此，本课题的共同科学问题可具体表述为：在给定设备能力、链路条件和负载阶段特征下，何时应采用请求级、阶段级或条件性任务内协同，且异构设备带来的有效工作量能否覆盖通信、调度、同步和恢复代价。研究实现对应两条路线：推理侧验证长上下文Prefill的HCP非对称协同；LLM-RL侧验证rollout、奖励评估和数据处理等阶段的异构承接与弹性执行。')
    set_text(p76, '宏观层面的成本与能源压力说明提高存量设备利用率具有现实意义，但本课题只对实际有效吞吐、资源占用和能效边界作实验性判断，不由宏观趋势推导固定加速结论。去中心化与多智能体方法的定位也随之明确：它们不是第三条研究线，而是当集中式或分层基线在跨域资源发现、局部恢复或局部观测决策上出现瓶颈时的对照机制；若引入后的决策收益不能覆盖训练、探索和恢复开销，则保留更简单的集中式或分层方案。')

    update_toc(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    main()
