from copy import deepcopy
from pathlib import Path

from docx import Document


SRC = Path(
    "/Users/stark_sim/Desktop/硕士课题/开题报告/"
    "开题报告_新版_沈达_面向大规模分布式AI的异构硬件合作可行性理论与关键技术研究.docx"
)
OUT = SRC.with_name(
    "开题报告_新版2_沈达_面向大模型训练与推理的异构算力协同可行性理论与关键技术研究.docx"
)
NEW_TITLE = "面向大模型训练与推理的异构算力协同可行性理论与关键技术研究"


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = deepcopy(para.runs[0]._r.rPr) if para.runs and para.runs[0]._r.rPr is not None else None
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def replace_all(doc, old, new):
    for para in doc.paragraphs:
        if old in para.text:
            set_text(para, para.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        set_text(para, para.text.replace(old, new))


def find_contains(doc, marker):
    for para in doc.paragraphs:
        if marker in para.text:
            return para
    raise ValueError(f"paragraph not found: {marker}")


def main():
    doc = Document(SRC)
    replace_all(doc, "面向大规模分布式AI的异构硬件合作可行性理论与关键技术研究", NEW_TITLE)
    replace_all(doc, "异构硬件合作", "异构算力协同")
    replace_all(doc, "异构合作", "异构算力协同")
    replace_all(doc, "大规模AI系统", "大规模大模型系统")
    replace_all(doc, "AI基础设施", "大模型基础设施")
    replace_all(doc, "其他分布式 AI 工作负载", "其他分布式大模型工作负载")
    replace_all(doc, "不同 AI 任务", "不同大模型训练与推理任务")
    replace_all(doc, "典型 AI 任务", "典型大模型训练与推理任务")

    source_para = find_contains(doc, "本课题即源于对上述问题的思考与前期实践")
    layout_para = find_contains(doc, "从研究布局来看")
    scale_para = find_contains(doc, "规模层面的矛盾同样突出")
    value_para = find_contains(doc, "就研究价值而言")
    significance_para = find_contains(doc, "这些工作的实际意义在于")
    set_text(
        source_para,
        "本课题即源于对上述问题的思考与前期实践。申请人围绕异构环境下基于环结构的上下文并行通信开展研究，已完成 hetero-cp-ringattn 原型系统的核心代码开发，并在小型混合集群上获得了初步加速效果。在此基础上，课题以分布式大模型训练与推理系统为主要对象，参考 Dynamo 的生产级推理编排与执行框架、HetRL 的异构大模型强化学习训练调度方法，以及 Prime RL 和 Prime DiLoCo 在异步训练、弹性设备网格和去中心化通信方面的实践，研究异构算力在资源池、工作流、任务和算子等不同层次上的协同机制。"
    )
    set_text(
        layout_para,
        "从研究布局来看，现有工作呈现明显的层次割裂。Dynamo 等工业系统主要关注大模型推理服务的资源编排、Prefill/Decode 分离、KV 传输和执行引擎组织；HetRL 关注包含多个模型、多个任务及复杂依赖的异构 LLM 强化学习训练调度；Prime RL 和 Prime DiLoCo 则进一步展示了异步强化学习、弹性设备网格、弱同步通信、节点动态加入退出和跨地域训练的工程可行性。这些工作分别覆盖了服务层、工作流层和通信基础设施层，但对于不同厂商设备在同一任务甚至同一计算层内如何协同执行，仍缺少统一的可行性刻画与跨层机制。"
    )
    set_text(
        scale_para,
        "规模层面的矛盾同样突出。现有异构系统普遍沿用集中式调度架构，隐含着全局拓扑可见、链路均匀可靠的假设；当节点规模扩大、链路质量参差不齐时，全局调度的开销急剧上升，不可直连节点对的通信问题也无从解决。Prime DiLoCo 通过 ElasticDeviceMesh、心跳失效检测、动态进程组调整、异步检查点和弱同步通信，说明节点可加入退出的去中心化大模型训练具有工程可行性；Prime RL 则将异步 rollout、训练、评估、FSDP/EP/CP 和 vLLM 执行引擎整合到统一工作流中。上述实践表明，强化学习工作流与异构通信基础设施在实施上具有天然关联，但仍需要针对不同设备能力和链路条件建立统一的协同决策机制。"
    )
    set_text(
        value_para,
        "就研究价值而言，本课题首先希望在理论上补齐上述缺口，建立适用于大模型推理、监督微调、强化学习后训练和去中心化训练等场景的异构算力协同可行性判定模型，为给定的设备组合、网络条件和任务图给出资源分配与效率边界。其次，课题将在技术层面打通从建模到系统的通路：在推理场景中以 hetero-cp-ringattn 为基础实现多代 GPU 对长上下文 Prefill 的细粒度协同，在强化学习训练场景中研究多模型、多任务和多阶段依赖的异构调度，并借助 Prime DiLoCo 式弹性通信思想和 Dynamo 式可插拔执行接口，把资源编排、工作流调度和任务内部协同组织在统一框架中。"
    )
    set_text(
        significance_para,
        "这些工作的实际意义在于直接服务于大模型基础设施的降本增效。若中低端存量 GPU、不同厂商加速器与跨地域分散算力能够以可预测的方式参与大模型训练和推理，系统就不必完全依赖单一型号的高端设备；同时，动态调度和弹性协同能够提高已有资源的利用率，降低大模型服务与后训练的硬件门槛和能源开销。"
    )
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
