from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE = Path('/Users/stark_sim/Desktop/硕士课题/开题报告')
SRC = BASE / '开题报告_新版12_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx'
OUT = BASE / '开题报告_新版17_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx'
IMG_DIR = Path('/Users/stark_sim/VSCodeProjects/hetero-cp-ringattn/output/imagegen')


def clone_run_format(para):
    if para.runs and para.runs[0]._r.rPr is not None:
        return deepcopy(para.runs[0]._r.rPr)
    return None


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = clone_run_format(para)
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def find_para(doc, marker):
    for para in doc.paragraphs:
        if marker in para.text:
            return para
    raise ValueError(f'paragraph not found: {marker}')


def replace_para(doc, marker, text):
    """Replace one known body paragraph while retaining its paragraph/run formatting."""
    para = find_para(doc, marker)
    set_text(para, text)
    return para


def make_para(doc, text, style=None, bold=False, align=None, size=None):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align is not None:
        para.alignment = align
    return para


def insert_after(anchor, element):
    anchor_node = anchor._tbl if hasattr(anchor, '_tbl') else anchor._p
    element_node = element._tbl if hasattr(element, '_tbl') else element._p
    anchor_node.addnext(element_node)
    return element


def add_caption(doc, text):
    para = make_para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_figure(doc, anchor, image_path, caption, width=6.1):
    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_before = Pt(3)
    image_para.paragraph_format.space_after = Pt(0)
    image_para.add_run().add_picture(str(image_path), width=Inches(width))
    insert_after(anchor, image_para)
    cap = add_caption(doc, caption)
    insert_after(image_para, cap)
    return cap


def set_cell(cell, text, bold=False, size=9.5):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    for margin in ('top', 'start', 'bottom', 'end'):
        pass


def add_table_after(doc, anchor, title, headers, rows, widths):
    cap = add_caption(doc, title)
    insert_after(anchor, cap)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Normal Table'
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, size=8.7)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    insert_after(cap, table)
    note = make_para(doc, '注：表中“新增机制”仅指本课题拟实现或验证的增量，不把参考框架已有能力写作创新。', size=9)
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(6)
    insert_after(table, note)
    return note


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    doc = Document(SRC)

    # The RL route is an engineering integration problem. The following
    # replacements keep the claims aligned with known prior art rather than
    # presenting stage scheduling or elasticity as a new RL algorithm.
    replace_para(
        doc,
        '本课题即源于对上述问题的思考与前期实践。',
        '本课题即源于对上述问题的思考与前期实践。申请人已完成hetero-cp-ringattn原型的核心开发，围绕非均匀上下文切片、K/V P2P ring和online softmax开展了跨设备验证。后续研究固定为两条互相独立、分别实施的研究路线；二者仅在设备画像和评价维度上相互借鉴：推理侧以长上下文Prefill为对象，在vLLM或Dynamo类服务接口下研究HCP准入、非对称切分和回退；LLM-RL侧以rollout、奖励计算、评估和数据处理等阶段为对象，整合HetRL的异构阶段放置、Prime RL的异步工作流与Prime DiLoCo的弹性恢复思路，形成面向本地异构设备池的后训练阶段承接原型。课题不把通用大模型预训练、跨弱链路TP/DP或HCP参与策略更新列为目标。',
    )
    replace_para(
        doc,
        '从研究布局来看，现有工作已经分别覆盖了若干可复用的层次。',
        '从研究布局来看，现有工作已经分别覆盖了若干可复用的层次。Dynamo、vLLM及相关系统提供请求编排、Prefill/Decode分离、KV管理和执行接口；HetRL关注LLM-RL阶段在异构GPU上的放置；Prime RL提供异步rollout—训练—评估工作流，Prime DiLoCo提供弹性设备网格、弱同步和检查点机制，StreamRL与AuroraRL进一步说明阶段解耦、弹性rollout和动态节点协作已有公开先例。因此，RL线不把“异构+异步+弹性”笼统描述为研究空白，而是以现有能力为组件，研究适配不同后端和设备能力的工程组合、运行约束及其收益边界。',
    )
    replace_para(
        doc,
        '本课题的研究价值不在于提出一个覆盖所有AI任务的异构大框架，',
        '本课题的研究价值不在于提出一个覆盖所有AI任务的异构大框架，而在于形成可执行的负载承接边界。对推理侧，研究HCP在长上下文Prefill中的容量感知非对称协同，并明确它与vLLM/Dynamo已有请求和阶段调度的分工；对LLM-RL侧，形成一个由现有阶段放置、异步流水线和弹性恢复机制组合而成的异构设备池承接原型，使rollout、奖励评估和数据处理可在满足版本、样本和恢复约束时接入不同设备，策略更新仍保留在高带宽同构或局部同构设备组。',
    )
    replace_para(
        doc,
        'Prime RL的公开工作流把异步rollout、训练、评估以及vLLM等执行引擎组织在同一后训练管线中，',
        'Prime RL的公开工作流把异步rollout、训练、评估以及vLLM等执行引擎组织在同一后训练管线中，证明异构设备可以通过阶段级和异步方式共同提供有效工作；Prime DiLoCo的ElasticDeviceMesh、异步检查点和节点弹性提供了动态加入退出与弱同步的工程参考。StreamRL和AuroraRL也分别给出阶段解耦、跨数据中心异构资源、弹性rollout actor与版本约束的公开先例。由此可见，广义的“异构+LLM-RL+动态弹性”并非空白；本课题不重复实现通用训练框架，而是把可复用能力收束为面向本地异构设备池的统一运行路径，验证设备能力、任务状态和故障条件共同约束下的实际承接价值。',
    )
    replace_para(
        doc,
        '综合来看，LLM-RL异构支持的研究缺口不是“再造一个RL算法”，',
        '综合来看，LLM-RL异构支持不需要再造一个RL算法，也不宜把已有框架的单项功能改写为“首创”。本课题的工程问题是：在一个由不同代际、不同厂商或不同后端设备构成的算力池中，如何把设备画像、阶段放置、异步样本流、版本/租约管理和节点恢复组合成一条可运行、可回放、可测量的后训练承接路径；并在吞吐、样本有效性、版本滞后和恢复成本之间给出可验证边界。',
    )
    replace_para(
        doc,
        '国内关于LLM-RL的公开研究和工程实践更多集中在RLHF放置、流水化和资源利用问题。',
        '国内关于LLM-RL的公开研究和工程实践更多集中在RLHF放置、流水化和资源利用问题。自适应RLHF放置工作将不同阶段映射到不同设备；RLBoost考察抢占式资源对rollout的承接；RolloutPipe研究rollout与训练的流水化重叠。这些工作说明生成、奖励和评估等阶段适合首先从工作流层利用异构资源，但也表明策略更新、版本滞后、样本有效性和恢复成本仍需受到约束。与本课题相关的HETHUB、AutoHet等异构训练框架可作为训练侧对照，不能直接替代LLM-RL阶段承接实验；本课题的工程组合也不宣称覆盖所有异构训练问题。',
    )
    replace_para(
        doc,
        '本课题以异构加速卡组合能够实际承接工作负载为目标，',
        '本课题以异构加速卡组合能够实际承接工作负载为目标，选择两类范围明确、可以分别开展实验的负载进行研究：一是长上下文大模型推理中的Prefill阶段，重点验证容量不同、算力不同的设备能否通过非对称上下文并行共同完成一次注意力计算；二是LLM-RL后训练中的rollout、奖励评估和数据处理等阶段，重点构建一个整合异构放置、异步样本流和弹性恢复的设备池承接原型，验证不同设备能否在不承担高频梯度同步的前提下共同支撑后训练工作流。两条路线分别建立设备、网络和负载特征画像，并采用可比的评价维度，但不追求把异构设备组成同构集群，也不把高通信的TP、DP或策略更新跨弱链路混合执行作为研究目标。',
    )
    replace_para(
        doc,
        '为保证研究内容可实施，两条路线分别以各自框架和原型作为基线与实现载体。',
        '为保证研究内容可实施，两条路线分别以各自框架和原型作为基线与实现载体。vLLM、Dynamo类框架负责推理服务、请求路由、Prefill/Decode分离和KV管理；HetRL、Prime RL、Prime DiLoCo、StreamRL与AuroraRL分别作为RL阶段放置、异步流水线、弹性同步、阶段解耦和动态节点协作的参考。推理线新增的是HCP的准入判断、非对称切分和服务适配；RL线的工作不是另造调度算法，而是以任务图和运行约束把已有能力组合为可复现的异构设备池承接方案。HCP不进入RL线。',
    )
    replace_para(
        doc,
        '3.1.2  LLM-RL阶段任务的异构承接与弹性执行',
        '3.1.2  异构设备池上的LLM-RL后训练阶段承接与弹性运行',
    )
    replace_para(
        doc,
        '本研究以LLM-RL后训练工作流为对象，重点承接rollout、奖励计算、评估和数据处理等阶段。',
        '本研究以LLM-RL后训练工作流为对象，拟形成异构设备池上的阶段承接与弹性运行原型。原型重点承接rollout、奖励计算、评估和数据处理等阶段：以可回放任务图串联prompt、trajectory、logprob、reward、mask和策略版本；以设备画像描述可用后端、模型/精度、显存、吞吐、链路和队列；以异步样本流把不同速度的worker接入同一训练回合；以版本租约和样本校验控制滞后及数据对应关系；以任务租约、检查点和重派发处理节点加入、退出与链路降级。策略更新限制在高带宽同构或局部同构设备组。该路线是对已有异构放置、异步流水线和弹性恢复能力的工程组合与边界验证，不重新设计RL算法，也不把HCP用于策略更新或梯度同步。',
    )
    replace_para(
        doc,
        'RL线的新增模块包括阶段能力合同、阶段放置与并发控制、版本时效和样本完整性检查，',
        'RL线的实现由五类相互配合的运行机制构成，而非五个彼此独立的“创新点”：设备画像与阶段放置决定哪些worker能够承接何种任务；异步样本流与队列控制吸收不同设备的速度差异；版本租约与样本完整性检查保证轨迹和策略状态可追溯；任务租约、检查点和重派发处理worker加入、失效和恢复；统一观测记录吞吐、滞后、无效样本和恢复代价。HetRL、Prime RL、Prime DiLoCo、StreamRL和AuroraRL分别提供可复用的先例或对照，本课题不把其已有功能写成本课题创新，而是验证在本地异构设备池中将这些能力组合运行的可行性与边界。',
    )
    replace_para(
        doc,
        '实验设置固定模型版本、token预算、奖励模型和任务集合，',
        '实验设置固定模型、token预算、奖励模型、任务集合与资源预算，比较固定同位部署、静态阶段放置、异步流水线、弹性设备池承接四类方案。前两类分别代表常规RLHF部署和HetRL式阶段放置，第三类代表Prime RL式工作流，第四类为本课题将设备画像、异步样本流、版本/任务租约与故障回退组合后的工程原型。主要观察有效rollout吞吐、单位有效样本成本、策略更新等待时间、版本滞后、样本有效率、资源利用率、通信量和任务质量；随后注入节点退出、链路降级和检查点恢复，测量重构时间、恢复后版本一致性和策略质量。只有异构设备带来的有效工作量增加能够覆盖异步、通信和恢复代价，才认为该工程组合具有实际价值。',
    )
    replace_para(
        doc,
        '在稳定资源池中，先以HetRL式启发式/ILP或约束优化作为可解释基线，',
        '在稳定资源池中，先以HetRL式启发式/ILP或约束优化作为可解释基线，比较固定同位部署、静态阶段放置与统一设备池承接；在动态资源池中，再与Prime RL式异步流水线、Prime DiLoCo式弹性设备网格、StreamRL式阶段解耦及AuroraRL式动态worker协作分别对照。研究结论只回答“在何种设备差异、阶段比例、版本约束和故障条件下，这种工程组合值得启用”，不把系统运行成功等同于新的优化算法或普适加速。',
    )
    replace_para(
        doc,
        '3.2.2  LLM-RL阶段承接与弹性执行方案',
        '3.2.2  异构设备池LLM-RL阶段承接与弹性运行方案',
    )
    replace_para(
        doc,
        '第一步，以可回放的LLM-RL阶段任务图实现研究基线，',
        '第一步，构造可回放的LLM-RL阶段任务图和设备画像，而非直接改写完整训练框架。任务图为rollout、奖励/评估、数据处理、策略更新和检查点记录模型副本、输入输出规模、计算/显存需求、参数版本和可容忍等待时间；每条样本记录prompt、trajectory、logprob、reward、mask和策略版本的对应关系。画像同时记录候选设备的后端能力、模型/精度兼容性、显存、吞吐、网络、队列和恢复窗口。该层可先接入公开工作流的轨迹或小规模开源管线，验证任务状态与调度逻辑，不把Prime RL或HetRL代码本身视为需要重复开发的对象。',
    )
    replace_para(
        doc,
        '第二步，实现阶段能力合同、放置与并发控制。',
        '第二步，实现面向异构设备池的承接控制面。能力合同首先排除无法运行目标模型或不满足显存/精度要求的设备；放置器以规则/启发式提供低开销在线决策，并以ILP/约束优化提供小规模对照；异步样本流按worker实测速率调节并发和队列；版本/任务租约在超出滞后阈值、样本字段不完整或预计网络时间过高时拒绝接入或重派发。调度器优先将可弹性扩展的rollout、奖励评估和数据处理放到匹配的异构资源，将高频更新约束在高带宽同构或局部同构设备组，并始终保留不迁移、降低并发度或暂停接入的回退选择。',
    )
    replace_para(
        doc,
        '第三步，验证异构RL支持的增量价值。',
        '第三步，验证工程组合的增量价值。设置固定同位部署、静态异构阶段放置、仅异步流水线和弹性设备池承接四种方案，分别对应常规RLHF部署、HetRL式阶段放置、Prime RL式异步工作流和本课题的组合原型。所有方案固定模型版本、token预算、奖励模型、任务集合与资源预算，比较rollout吞吐、单位有效样本成本、策略更新等待时间、版本滞后、样本有效率、资源利用率、网络传输量和任务质量；通过组件消融区分设备画像、异步样本流、版本/任务租约和回退机制各自带来的影响。',
    )
    replace_para(
        doc,
        '第四步，在可控故障注入与弱连接场景验证弹性边界。',
        '第四步，在可控故障注入与弱连接场景验证弹性运行边界。对节点加入退出、链路降级、检查点恢复和worker失效分别记录重构时间、有效吞吐、参数/样本版本一致性和恢复后策略质量。Prime DiLoCo的ElasticDeviceMesh、异步检查点和live recovery，以及AuroraRL的动态actor和版本约束作为参考；本课题只研究在何种恢复窗口、队列状态和阶段依赖下允许节点进入、暂停或退出这一工程工作流。若组合机制没有显著减少停机或提高有效资源利用率，便保留集中式/分层方案，而不以“接入更多节点”代替真实收益。',
    )
    replace_para(
        doc,
        '（3）RL目标：完成rollout、奖励评估和数据处理阶段的异构放置、',
        '（3）RL目标：完成面向异构设备池的LLM-RL后训练阶段承接原型，集成rollout、奖励评估和数据处理的异构放置、异步样本流、版本/任务租约、节点失效恢复和局部回退。',
    )

    # Add macro-to-specific transitions without changing the established scope.
    p = find_para(doc, '本课题即源于对上述问题的思考与前期实践。')
    new = make_para(doc, '从研究对象的可拆分性出发，本课题将“异构承接”理解为一组带条件的系统选择：先识别负载阶段和状态依赖，再估计通信、显存、排队与恢复成本，最后决定采用请求级、阶段级或任务内协同。这样的定义既能容纳真实环境中的设备差异，也能避免将所有连接在同一网络中的加速卡都视为可以共同执行任意任务。')
    insert_after(p, new)
    p = find_para(doc, '这些工作的实际意义在于让异构资源以可预测的方式承接真实工作')
    new = make_para(doc, '因此，课题的结论将同时包含正向结果和负向边界：正向结果说明某类设备组合、上下文规模或阶段比例下存在可复现的净收益；负向边界说明通信、拖尾、版本滞后或恢复开销超过收益时应当拒绝协同。将两类结果同时写入研究结论，能够使后续系统设计具有明确的使用条件。')
    insert_after(p, new)

    p = find_para(doc, '本章按“服务组织—任务内协同—后训练工作流—弱连接协作”的顺序梳理国外研究')
    new = make_para(doc, '阅读本章时需要区分三种证据。第一种是已经在工业或开源系统中稳定存在的工程能力，例如请求路由、阶段分离、KV管理和异步流水线；第二种是论文中在受控环境下验证的并行机制，例如非均匀序列切分和低频参数同步；第三种是本课题尚需通过原型和实验确认的组合能力，即在明确通信和恢复约束下，让异构设备承接一部分真实工作。后文将按这一层次标记研究空缺，避免把“可借鉴”误写成“已解决”。')
    insert_after(p, new)
    p = find_para(doc, '由此，本课题的两条路线可共享的分析问题可具体表述为')
    new = make_para(doc, '这一综述还给出一个重要的节奏判断：研究不应从“如何把更多设备加入任务”开始，而应从“不加入会损失什么、加入后新增什么代价”开始。推理线先验证长上下文是否存在显存墙，再决定是否打开HCP；RL线先确认阶段边界和版本约束，再决定是否把弹性设备纳入工作流。两条路线均以能够被实验拒绝的准入条件作为起点。')
    insert_after(p, new)

    # Add a domestic-evidence subsection before the synthesis, keeping positive
    # capability claims paired with explicit communication and deployment limits.
    p = find_para(doc, '国内关于LLM-RL的公开研究和工程实践更多集中在RLHF放置、流水化和资源利用问题。')
    domestic = make_para(doc, '2.2.3  昇腾集群与国产异构软件栈的工程证据', bold=True)
    insert_after(p, domestic)
    p = domestic
    for text in [
        '昇腾相关公开成果为国内异构算力研究提供了比“设备型号对比”更具体的证据。PanGu-Σ在昇腾910集群和MindSpore上完成了万亿参数稀疏语言模型训练，通过专家计算与存储分离扩大了可用的计算与存储边界[40]；该工作说明国产NPU集群能够承载大模型训练，但其收益来自模型结构、并行方式与系统软件的协同设计，并不意味着任意异构卡都可以直接混合执行高频同步训练。',
        'CloudMatrix384进一步展示了384个昇腾910 NPU与192个鲲鹏CPU通过统一总线构成的超节点，并在推理侧采用Prefill、Decode和缓存的独立扩展[41]。该案例对本课题有两点启示：一是国产集群已经把异构资源池化和阶段解耦推进到生产级系统；二是其MoE专家并行和KV访问依赖专用高带宽全互联，反向说明普通网络下的跨厂商任务内协同必须谨慎设定边界。',
        '从软件栈看，MindSpore/MindSpeed提供了数据并行、张量并行、流水并行和混合并行等分布式能力[42]，vLLM-Ascend则通过插件方式把昇腾后端接入主流推理接口[44]。这些工程能力可作为本课题的适配基线，但“后端可运行”与“跨设备共同完成一次任务”是两个不同层次的问题：前者解决算子和运行时兼容，后者还需要设备画像、通信预算、准入和失败回退。',
        '国内异构协同的学术证据也不只来自大模型。ReDSEa在鲲鹏CPU与昇腾910组成的超算异构系统上对三角方程求解进行自动映射、负载均衡和调度[43]，证明了计算密集型任务可以通过性能模型获得异构收益；但其任务具有明确的分块结构，与LLM推理或RL后训练不同。因此，本课题将昇腾集群和ReDSEa作为“异构承接可行但依赖负载结构”的国内证据，而不是直接复用其性能结论。'
    ]:
        q = make_para(doc, text)
        insert_after(p, q)
        p = q
    framing = make_para(doc, '图4将异构现状、已有系统的承接粒度和本课题拟补足的研究空档放在同一比较轴上；昇腾与CloudMatrix384作为国内工程证据嵌入主流承接方式，而不单独代表本课题的研究对象。')
    insert_after(p, framing)
    p = framing
    state = IMG_DIR / 'heterogeneous-service-landscape-v1.png'
    if state.exists():
        cap = add_figure(doc, p, state, '图4  异构算力服务现状、主流承接方式与研究空档', width=6.0)
        p = cap
    rows = [
        ('Dynamo / vLLM / llm-d', '请求路由、Prefill/Decode分离、KV管理、后端接口', '适合请求级/阶段级承接；不解决单请求任务内HCP'),
        ('CloudMatrix384 / vLLM-Ascend', '昇腾+鲲鹏专用互联；国产后端插件与阶段解耦', '证明国产异构池化可落地；依赖专用高带宽互联与软件栈'),
        ('HetRL / Prime RL / StreamRL / AuroraRL', '异构阶段放置、异步rollout、阶段解耦、弹性worker与版本约束', '说明RL阶段承接已有先例；不宜把广义异构弹性写作空白'),
        ('本课题', '长上下文Prefill的HCP；RL设备池的组合式承接原型', '以画像、样本流、租约和回退把已有能力组合为可测工程路径'),
    ]
    note = add_table_after(doc, p, '表3  主流异构承接方式与本课题增量边界', ['代表工作/对象', '已展示能力', '对本课题的启示与限制'], rows, [1.6, 2.8, 2.0])

    # Add overview figure and a compact research-plan matrix at the beginning of Chapter 3.
    p = find_para(doc, '两条路线均从相同的分析维度出发')
    overview = IMG_DIR / 'heterogeneous-two-research-directions-overview-v4-labeled.png'
    if overview.exists():
        cap = add_figure(doc, p, overview, '图1  两条独立研究路线及其异构负载承接边界', width=6.0)
        p = cap
    rows = [
        ('长上下文Prefill', '输入长度、KV规模、显存、P2P带宽/时延、TTFT约束', 'HCP准入；capacity-aware非对称seq_chunk_len/block_size；K/V P2P ring；超时和容量变化回退', '单设备、同构CP、请求/PD路由、固定均分HCP；TTFT、尾延迟、显存峰值、网络字节数、回退率'),
        ('LLM-RL后训练阶段', '阶段输入输出、后端兼容性、版本、队列、恢复窗口', '设备画像与阶段放置；异步样本流；版本/任务租约；节点失效重派发', '固定同位、静态阶段放置、异步流水线、弹性设备池承接；有效样本吞吐、样本有效率、版本滞后、恢复时间'),
    ]
    note = add_table_after(doc, p, '表1  两条研究路线的方案闭环', ['研究路线', '输入与约束', '本课题新增机制', '基线与主要评价'], rows, [1.0, 1.7, 2.1, 1.6])

    p = find_para(doc, '3.2  实施方案')
    value_note = make_para(doc, '图5不再重复研究路线和实验矩阵，而是区分参考框架已经提供的能力、本课题拟完成的工程组合，以及仍需实验验证的运行边界。推理与LLM-RL在图中保持为上下两条独立路线。')
    insert_after(p, value_note)
    matrix = IMG_DIR / 'heterogeneous-value-boundary-v1.png'
    if matrix.exists():
        add_figure(doc, value_note, matrix, '图5  现有异构系统能力、本课题工程组合与验证边界', width=6.0)

    # Route-specific figures stay inside their own sections.
    p = find_para(doc, '本研究只把长上下文Prefill作为任务内异构协同的验证对象')
    hcp = IMG_DIR / 'heterogeneous-hcp-prefill-detail-v3.png'
    if hcp.exists():
        add_figure(doc, p, hcp, '图2  长上下文Prefill中HCP的非对称切分、K/V环传递与准入回退', width=5.9)
    p = find_para(doc, '本研究以LLM-RL后训练工作流为对象')
    rl = IMG_DIR / 'heterogeneous-llm-rl-stage-admission-v3.png'
    if rl.exists():
        add_figure(doc, p, rl, '图3  异构设备池上的LLM-RL阶段承接与弹性运行', width=5.9)

    # Add a capability boundary table in the completed-work chapter.
    p = find_para(doc, '（1）形成 hetero-cp-ringattn 开源代码库和异构通信原型')
    rows = [
        ('vLLM / Dynamo类框架', 'API、请求路由、Prefill/Decode分离、KV生命周期、常规worker执行', '不重写服务框架；增加HCP候选计划的调用接口与回退状态'),
        ('HCP原型', '非均匀seq_chunk_len、block_size、K/V P2P ring、online softmax协议基础', '补充设备/链路画像、准入代价估计、容量感知切分和服务适配验证'),
        ('HetRL / Prime RL / StreamRL', '阶段放置、异步rollout、训练/评估工作流、阶段解耦', '以可回放任务图和异步样本流将既有能力适配到异构设备池'),
        ('Prime DiLoCo / AuroraRL', '弹性设备网格、检查点恢复、动态worker与版本约束', '以版本/任务租约和重派发验证节点加入/退出、链路降级和恢复窗口'),
    ]
    add_table_after(doc, p, '表4  参考框架已有能力与本课题新增机制', ['参考对象', '已有能力（作为基线或接口）', '本课题新增或待验证内容'], rows, [1.3, 3.0, 2.1])

    # Make key section headings visually distinct while preserving document fonts.
    for marker in ['1.  课题来源及研究的背景和意义', '2. 国内外在该方向的研究现状及分析', '3. 主要研究内容及实施方案', '4. 预 期 达 到 的 目 标', '5. 已 完 成 的 研 究 工 作 与 进 度 安 排', '6.  为完成课题已具备和所需的条件和经费', '7.  预计研究过程中可能遇到的困难和问题，以及解决的措施', '8.  主要参考文献']:
        try:
            para = find_para(doc, marker)
            if para.runs:
                para.runs[0].bold = True
        except ValueError:
            pass

    # Extend the bibliography with the domestic evidence cited in section 2.2.
    p = find_para(doc, 'Prime Intellect. PRIME-DiLoCo: Elastic and Low-Communication Distributed Training')
    refs = [
        '[40] Ren X, Zhou P, Meng X, et al. PanGu-Σ: Towards Trillion Parameter Language Model with Sparse Heterogeneous Computing[EB/OL]. arXiv:2303.10845, 2023.',
        '[41] Zuo P, Lin H, Deng J, et al. Serving Large Language Models on Huawei CloudMatrix384[EB/OL]. arXiv:2506.12708, 2025.',
        '[42] MindSpore. Distributed Parallelism Overview[EB/OL]. https://www.mindspore.cn/tutorials/experts/en/r2.3.1/parallel/overview.html, 2025.',
        '[43] Zacharopoulos G, Bournias I, Vlacic V, et al. ReDSEa: Automated Acceleration of Triangular Solver on Supercloud Heterogeneous Systems[EB/OL]. arXiv:2305.19917, 2023.',
        '[44] vLLM Project. vLLM Ascend: Community-Maintained Hardware Plugin for Running vLLM on Ascend NPU[EB/OL]. https://github.com/vllm-project/vllm-ascend, 2025.',
        '[45] StreamRL. Disaggregating LLM Reinforcement Learning: Efficient Heterogeneous Resource Utilization for Large-Scale RL Training[EB/OL]. arXiv:2504.15930, 2025.',
        '[46] AuroraRL. Decentralized Elastic Reinforcement Learning for LLMs in Heterogeneous Environments[EB/OL]. arXiv:2602.11456, 2026.',
    ]
    for ref in refs:
        q = make_para(doc, ref, size=9)
        insert_after(p, q)
        p = q

    doc.save(OUT)
    print(f'wrote {OUT}')


if __name__ == '__main__':
    main()
