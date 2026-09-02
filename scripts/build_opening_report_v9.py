from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path('/Users/stark_sim/Desktop/硕士课题/开题报告/开题报告_新版8_沈达_面向大模型推理与强化学习后训练的异构算力协同可行性理论与关键技术研究.docx')
OUT = SRC.with_name('开题报告_新版9_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx')
TITLE = '面向大模型推理与强化学习后训练的异构算力负载承接研究'

def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = deepcopy(para.runs[0]._r.rPr) if para.runs and para.runs[0]._r.rPr is not None else None
    para.clear()
    run = para.add_run(text)
    if ppr is not None: para._p.insert(0, ppr)
    if rpr is not None: run._r.insert(0, rpr)

def remove_para(para):
    para._element.getparent().remove(para._element)

def first(doc, pred):
    return next(p for p in doc.paragraphs if pred(p.text))

def replace_all(doc, old, new):
    for p in doc.paragraphs:
        if old in p.text: set_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text: set_text(p, p.text.replace(old, new))

def update_toc(doc):
    replacements = {
        '3.1.1  异构推理的多层次协同扩展': '3.1.1  长上下文Prefill的异构非对称协同',
        '3.1.2  强化学习工作负载的异构承接': '3.1.2  LLM-RL阶段任务的异构承接与弹性执行',
        '3.2.1  异构推理协同扩展方案': '3.2.1  长上下文Prefill异构协同实施方案',
        '3.2.2  强化学习负载弹性编排方案': '3.2.2  LLM-RL阶段承接与弹性执行方案',
    }
    for p in doc.part._element.xpath('.//w:sdt//w:p'):
        text = ''.join(t.text or '' for t in p.iter(qn('w:t')))
        for old, new in replacements.items():
            if old in text:
                for t in p.iter(qn('w:t')):
                    if t.text and old in t.text: t.text = t.text.replace(old, new)
                break

def main():
    doc = Document(SRC)
    replace_all(doc, '面向大模型推理与强化学习后训练的异构算力协同可行性理论与关键技术研究', TITLE)
    replace_all(doc, '面向大模型训练与推理的异构算力协同可行性理论与关键技术研究', TITLE)
    replace_all(doc, '3. 主要研究内容及研究方案', '3. 主要研究内容及实施方案')
    set_text(first(doc, lambda t: t.strip() == '3.1  主要研究内容'), '3.1  研究内容')
    set_text(first(doc, lambda t: t.strip() == '3.2  研究方案'), '3.2  实施方案')

    set_text(first(doc, lambda t: t.startswith('本课题围绕异构算力协同的可行性与可持续性')),
        '本课题以异构加速卡组合能够实际承接工作负载为目标，选择两类范围明确、可以分别开展实验的负载进行研究：一是长上下文大模型推理中的Prefill阶段，重点验证容量不同、算力不同的设备能否通过非对称上下文并行共同完成一次注意力计算；二是LLM-RL后训练中的rollout、奖励评估和数据处理等阶段，重点验证不同设备能否在不承担高频梯度同步的前提下共同支撑后训练工作流。两条路线共享设备、网络和负载特征画像，但不追求把异构设备组成同构集群，也不把高通信的TP、DP或策略更新跨弱链路混合执行作为研究目标。')
    set_text(first(doc, lambda t: t.startswith('本课题将“协同粒度选择”作为贯穿两点的核心问题')),
        '两条路线采用同一条判断原则：先根据设备算力和显存、后端兼容性、链路带宽与时延、状态传递频率以及任务SLO，判断异构卡适合承接请求、阶段还是任务内的一部分；只有预计收益能够覆盖通信、调度、同步和恢复开销时才扩大协同范围，否则回退到单设备、同构设备组或更粗粒度的阶段编排。研究成果以可执行的放置规则、切分计划和回退条件为主，不预设所有异构组合都能获得加速。')
    set_text(first(doc, lambda t: t.startswith('为使研究内容能够直接落到实现与实验')),
        '为保证研究内容可实施，现有框架和原型只作为基线与承载接口。vLLM、Dynamo类框架负责推理服务、请求路由、Prefill/Decode分离和KV管理；HetRL、Prime RL和Prime DiLoCo分别作为RL阶段放置、异步流水线和弹性同步的参考实现。推理线新增的是HCP的准入判断、非对称切分和服务适配；RL线新增的是阶段能力合同、版本/队列约束、异构阶段放置和节点失效后的回退。HCP不进入RL线。')
    set_text(first(doc, lambda t: t.startswith('研究的评价不仅报告平均吞吐提升')),
        '评价以“异构卡是否真正承接了有效工作”为中心：推理侧同时记录TTFT、TPOT、尾延迟、显存峰值、KV传输量、计划耗时和回退率；RL侧同时记录有效rollout吞吐、单位有效样本成本、策略更新等待时间、版本滞后、样本有效率、资源利用率和恢复时间。正结果用于确定可用条件，负结果用于界定通信、上下文长度、负载比例和异构度的失效边界。去中心化和多智能体方法不列为独立研究点，只在集中式或分层基线确有瓶颈时作小规模对照。')

    set_text(first(doc, lambda t: t.startswith('3.1.1  ')), '3.1.1  长上下文Prefill的异构非对称协同')
    set_text(first(doc, lambda t: t.startswith('3.2.1  ')), '3.2.1  长上下文Prefill异构协同实施方案')
    set_text(first(doc, lambda t: t.startswith('面向大模型在线推理')),
        '本研究只把长上下文Prefill作为任务内异构协同的验证对象。首先测量候选设备在不同上下文长度和batch下的Prefill速率、可用显存、attention瞬态工作区以及P2P带宽和时延；其次以单设备、同构CP和异构请求/PD路由为基线，判断当前请求是否存在显存墙或TTFT瓶颈；最后在满足通信预算时调用HCP，把输入序列按设备容量和attention速率切成不等长区间，由各domain保留本地Q并沿P2P ring传递K/V块，在线合并结果。研究目标是得到一套可执行的准入规则、seq_chunk_len与block_size生成方法和超时回退路径，而不是重做vLLM或Dynamo的通用调度器。')
    set_text(first(doc, lambda t: t.startswith('HCP在该基线之下提供不同层次的能力')),
        '现有HCP原型已经具备非均匀seq_chunk_len、block_size、K/V P2P ring和online softmax等协议基础。本课题要补充的具体工作包括：建立设备与链路画像；实现对单设备、请求/PD路由和HCP路径的代价估计；自动生成容量—通信感知的非对称切分；通过vLLM或Dynamo类接口下发执行计划；在超时、容量变化、数值校验失败或SLO风险出现时回退。HCP只服务于被批准的长上下文Prefill请求，不替代TP、DP或MoE的高频同步。')
    set_text(first(doc, lambda t: t.startswith('本课题在推理线的新增工作由三部分组成')),
        '实验上固定模型、请求集和网络条件，依次比较单设备、同构CP、静态异构请求/PD路由、固定均分HCP和非对称HCP。改变上下文长度、显存差异、计算速率比、带宽/时延和并发度，分别记录正确性、TTFT、吞吐、尾延迟、显存峰值、网络字节数、计划耗时和回退率。只有当非对称HCP在显存可承接性或端到端指标上超过相应基线时，才把该配置计为有效异构承接；否则记录为不准入条件。')

    set_text(first(doc, lambda t: t.startswith('3.1.2  ')), '3.1.2  LLM-RL阶段任务的异构承接与弹性执行')
    set_text(first(doc, lambda t: t.startswith('3.2.2  ')), '3.2.2  LLM-RL阶段承接与弹性执行方案')
    set_text(first(doc, lambda t: t.startswith('面向强化学习相关工作负载，自适应RLHF放置')),
        '本研究以LLM-RL后训练工作流为对象，重点承接rollout、奖励计算、评估和数据处理等阶段。首先把每个阶段表示为带有输入输出规模、显存/计算需求、后端类型、模型版本、网络依赖和可容忍等待时间的任务；其次根据设备能力合同，把生成类和评估类阶段放到可用的异构设备，把策略更新限制在高带宽同构或局部同构设备组；最后在节点加入、退出、链路降级和检查点恢复时，验证阶段任务能否暂停、迁移、降并发或回退。该路线研究的是异构卡对RL工作流的阶段承接，不是重新设计RL算法，也不把HCP用于策略更新或梯度同步。')
    set_text(first(doc, lambda t: t.startswith('RL线的研究对象与HCP严格解耦')),
        'RL线的新增模块包括阶段能力合同、阶段放置与并发控制、版本时效和样本完整性检查，以及worker加入/失效/恢复时的回退规则。调度器在放置前检查后端可用性、显存、队列长度、预计网络时间和恢复窗口；在执行中限制版本滞后和无效样本比例；在约束无法满足时降低并发度或停止迁移。Prime RL的异步工作流、HetRL的异构阶段优化和Prime DiLoCo的弹性设备网格分别作为对照，不把它们已经具备的功能写成本课题创新。')
    set_text(first(doc, lambda t: t.startswith('本课题在RL线的候选增量贡献')),
        '实验设置固定模型版本、token预算、奖励模型和任务集合，比较固定同位部署、静态异构阶段放置、仅异步流水线和能力合同驱动的动态阶段承接。主要观察有效rollout吞吐、单位有效样本成本、策略更新等待时间、版本滞后、样本有效率、资源利用率、通信量和任务质量。随后注入节点退出、链路降级和检查点恢复，测量重构时间、恢复后版本一致性和策略质量。只有异构设备带来的有效工作量增加能够覆盖异步、通信和恢复代价，才认为组合具有实际价值。')

    for p in list(doc.paragraphs):
        if p.text.startswith('RL线与推理线严格分工：HCP只属于推理线') or p.text.startswith('在稳定资源池上，先实现集中式或分层协调') or p.text.startswith('在弹性扩展阶段，模拟或实测节点加入退出'):
            remove_para(p)

    set_text(first(doc, lambda t: t.startswith('（1）理论目标：建立设备—网络—任务统一')), '（1）方法目标：形成设备—网络—负载特征表和协同粒度选择规则，明确异构卡适合承接请求、Prefill阶段、RL阶段任务还是不适合承接。')
    set_text(first(doc, lambda t: t.startswith('（2）方法目标：推理线形成容量感知')), '（2）推理目标：完成长上下文Prefill的HCP准入、非对称seq_chunk_len/block_size计划、vLLM或Dynamo类接口适配和回退机制。')
    set_text(first(doc, lambda t: t.startswith('（3）系统目标：实现可插拔')), '（3）RL目标：完成rollout、奖励评估和数据处理阶段的异构放置、版本/队列控制、节点失效恢复和局部回退原型。')
    set_text(first(doc, lambda t: t.startswith('（4）实验目标：在多代、多厂商')), '（4）实验目标：在小规模混合设备上分别验证推理侧HCP和RL侧阶段承接的正确性、有效吞吐、资源占用、通信代价、失效边界和恢复能力，不预设固定加速倍数。')
    set_text(first(doc, lambda t: t.startswith('（5）成果目标：围绕协同粒度')), '（5）成果目标：形成异构设备画像、推理HCP计划器、RL阶段承接原型、可复现实验基准以及相应论文/专利材料。')

    update_toc(doc)
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    main()
