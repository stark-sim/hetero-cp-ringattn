from copy import deepcopy
from pathlib import Path

from docx import Document


SRC = Path('/Users/stark_sim/Desktop/硕士课题/开题报告/开题报告_新版11_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx')
OUT = SRC.with_name('开题报告_新版12_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx')


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = deepcopy(para.runs[0]._r.rPr) if para.runs and para.runs[0]._r.rPr is not None else None
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def replace_all(doc, old, new):
    changed = 0
    for para in doc.paragraphs:
        if old in para.text:
            set_text(para, para.text.replace(old, new))
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        set_text(para, para.text.replace(old, new))
                        changed += 1
    return changed


def main():
    doc = Document(SRC)
    replacements = [
        ('两条互相独立但共享画像方法的路线', '两条互相独立、分别实施的研究路线；二者仅在设备画像和评价维度上相互借鉴'),
        ('两条路线共享设备、网络和负载特征画像', '两条路线分别建立设备、网络和负载特征画像，并采用可比的评价维度'),
        ('两条路线采用同一条判断原则', '两条路线均从相同的分析维度出发'),
        ('现有框架和原型只作为基线与承载接口', '两条路线分别以各自框架和原型作为基线与实现载体'),
        ('共同科学问题可具体表述为', '两条路线可共享的分析问题可具体表述为'),
        ('本课题的组合候选', '本课题的动态阶段承接方案'),
        ('这正是现有原型和现有服务框架尚未共同提供的新增机制', '这是推理线需要新增的机制，不能外推为RL线的共同模块'),
        ('为后续研究统一设备画像、计算—通信联合调度和多平台后端适配提供了可运行的系统基础', '为推理线后续开展设备画像、计算—通信联合调度和多平台后端适配提供了可运行的原型基础'),
        ('后续将把这些工程结果统一纳入协同粒度可行性模型', '后续将把设备与链路测量方法分别用于两条路线的协同粒度分析，其中HCP原型结果仅用于推理线'),
        ('否则以更简单的基线为最终系统方案', '否则以更简单的基线作为最终实现方案'),
        ('建立统一任务—资源—网络模型', '分别建立推理与RL负载的任务—资源—网络描述'),
        ('整合理论、算法和系统，完成综合实验、毕业论文撰写与答辩', '分别完成两条路线的实验总结，形成统一论文框架并完成毕业论文撰写与答辩'),
    ]
    changed = sum(replace_all(doc, old, new) for old, new in replacements)
    doc.save(OUT)
    print(f'wrote {OUT}; replacements={changed}')


if __name__ == '__main__':
    main()
