from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement


SRC = Path(
    "/Users/stark_sim/Desktop/硕士课题/开题报告/"
    "开题报告_初版_沈达_面向大规模分布式AI的异构硬件合作可行性理论与关键技术研究.docx"
)
OUT = SRC.with_name(
    "开题报告_新版_沈达_面向大规模分布式AI的异构硬件合作可行性理论与关键技术研究.docx"
)


def find_para(doc, prefix):
    for para in doc.paragraphs:
        if para.text.strip().startswith(prefix):
            return para
    raise ValueError(f"paragraph not found: {prefix}")


def copy_format(src, dst):
    if src._p.pPr is not None:
        dst._p.insert(0, deepcopy(src._p.pPr))
    if src.runs and src.runs[0]._r.rPr is not None and dst.runs:
        dst.runs[0]._r.insert(0, deepcopy(src.runs[0]._r.rPr))


def add_para(doc, elements, text, exemplar, page_break=False):
    para = doc.add_paragraph()
    run = para.add_run(text or "")
    copy_format(exemplar, para)
    if page_break:
        run.add_break(WD_BREAK.PAGE)
    elements.append(para._p)


def add_table(doc, elements, headers, rows, style=None):
    table = doc.add_table(rows=1, cols=len(headers))
    if style:
        table.style = style
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    elements.append(table._tbl)


def main():
    doc = Document(SRC)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "true")
    doc.settings.element.append(update_fields)
    section3 = find_para(doc, "3.")
    section6 = find_para(doc, "6.")
    section4 = find_para(doc, "4.")
    body_exemplar = find_para(doc, "面对异构硬件能力差异大")
    heading_exemplar = section3
    subheading_exemplar = find_para(doc, "3.1  ")
    subsubheading_exemplar = find_para(doc, "3.1.1")

    body = doc._element.body
    children = list(body)
    start = children.index(section3._p)
    end = children.index(section6._p)
    for child in children[start:end]:
        body.remove(child)

    elements = []

    add_para(doc, elements, "", body_exemplar, page_break=True)
    add_para(doc, elements, "3. 主要研究内容及研究方案", heading_exemplar)
    add_para(doc, elements, "3.1  主要研究内容", subheading_exemplar)
    add_para(
        doc,
        elements,
        "本课题以异构硬件合作的可行性与可持续性为核心问题，不把研究限定在某一种模型、某一种加速器或某一个并行算法上。研究内容按照“统一抽象—协同机制—动态适应—场景验证”的层次展开：首先建立设备、任务和网络的统一描述；随后研究异构资源的任务划分与计算通信协同；在此基础上研究弱连接环境下的去中心化组织和容错；最后利用多智能体强化学习等方法实现对动态状态的在线适配。长上下文大模型 Prefill、分布式训练和去中心化强化学习工作流作为具有代表性的验证场景，根据实验条件和阶段性结果逐步展开。",
        body_exemplar,
    )
    add_para(doc, elements, "3.1.1  异构合作可行性的统一理论建模与判定", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "研究面向异构合作的统一性能与约束模型。对设备的有效计算速率、显存容量、内存带宽、能耗和可靠性进行画像，对设备间链路的带宽、延迟、丢包率和可达性进行刻画，并将目标任务表示为包含计算节点、数据依赖和同步关系的任务图。在此基础上，研究给定设备集合完成给定任务的可行性判定、资源分配约束和效率上界，分析设备数量、能力差异、通信质量与任务结构之间的关系。该模型既用于推理任务，也为训练任务和其他分布式 AI 工作负载提供可迁移的描述方式。",
        body_exemplar,
    )
    add_para(doc, elements, "3.1.2  异构任务划分与计算—通信联合调度", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "研究根据设备能力进行非均匀任务划分的方法，重点解决计算负载、显存占用、通信代价和同步等待之间的耦合关系。针对长上下文 Prefill，研究序列块和 KV 块的容量感知切分、Ring Attention 中的环序与块大小选择，以及计算和通信的流水化重叠；针对分布式训练和其他任务，研究数据、流水线阶段、参数更新或子任务的非对称分配。通过统一的代价模型比较静态启发式、约束优化和在线调度方法，明确不同方法在不同异构程度和网络条件下的适用边界。",
        body_exemplar,
    )
    add_para(doc, elements, "3.1.3  面向弱连接环境的去中心化协同与容错通信", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "研究在全局拓扑不可见、节点能力不一致、链路质量动态变化的条件下组织协同计算的方法。主要包括局部邻居发现与能力交换、基于局部信息的 P2P 拓扑和逻辑环构建、中继路径选择、节点加入与退出、链路异常和设备失效时的局部重构，以及协同状态的一致性维护。研究目标不是简单替换中心式调度器，而是分析局部信息带来的性能损失与扩展收益，建立通信可靠性、任务连续性和重构开销之间的权衡机制。",
        body_exemplar,
    )
    add_para(doc, elements, "3.1.4  基于多智能体强化学习的动态自适应与能效优化", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "研究适用于动态异构集群的多智能体强化学习方法。将设备或协作域视为智能体，以自身负载、显存余量、邻居状态、链路质量、任务进度和历史重构开销作为局部观测，以任务份额调整、邻居选择、路径切换、并发度调节和节点接入后的资源重分配作为动作，设计同时考虑完成时间、吞吐、尾延迟、能耗和稳定性的奖励函数。比较集中训练—分散执行、值函数分解、Actor-Critic 及其轻量化变体，并研究强化学习与启发式优化、约束求解相结合的混合策略，使具体算法可以根据场景和实验结果灵活调整。",
        body_exemplar,
    )
    add_para(doc, elements, "3.1.5  多场景验证与适用边界分析", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "选取长上下文大模型 Prefill 作为首个端到端验证场景，以 hetero-cp-ringattn 为系统基础，验证异构 Ring Attention、KV 传输和容量感知切分；根据理论模型和系统成熟度，进一步在分布式强化学习、异构训练或去中心化算力协作场景中开展扩展实验。不同场景使用统一的设备画像和代价指标，比较单设备、同构并行、静态异构调度和动态协同方法，重点报告任务完成时间、吞吐、显存峰值、通信开销、能耗、扩展效率和故障恢复时间，并明确方法失效的条件。",
        body_exemplar,
    )

    add_para(doc, elements, "3.2  研究方案", subheading_exemplar)
    add_para(doc, elements, "3.2.1  统一模型与可行性判定方案", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "首先通过微基准测试建立设备和链路性能档案，将异构集群抽象为带属性的图 G=(V,E)，其中节点属性包括有效计算速率、显存容量、内存带宽、功耗和可靠性，边属性包括带宽、延迟、丢包率和可达性。其次，将不同 AI 任务表示为带资源需求和通信依赖的任务图，定义设备映射、数据分片、通信路径和同步点等决策变量。最后构造带有显存、时延、吞吐、能耗和可靠性约束的可行性判定问题，结合排队模型、关键路径分析和上下界估计，输出可行设备组合、资源分配建议以及不可行配置的原因。",
        body_exemplar,
    )
    add_para(doc, elements, "3.2.2  异构任务划分与计算—通信协同方案", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "在统一模型基础上设计分层调度器。上层根据任务类型、SLO 和设备可用资源筛选候选协作集合；中层联合决定任务分片比例、设备映射、通信顺序和流水化参数；底层通过可插拔通信后端执行 KV 块或其他中间数据的发送、接收和转发。对于 Ring Attention，重点实现 capacity-aware 的非均匀序列分片、KV ring 和计算通信重叠，并保留可靠的字节流回退路径；对于训练或其他任务，复用同一套资源画像和代价接口，替换具体任务图与调度策略。",
        body_exemplar,
    )
    add_para(doc, elements, "3.2.3  去中心化拓扑与容错通信方案", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "每个节点通过周期性探测维护有限规模的邻居视图，并交换必要的设备能力、任务状态和链路统计信息。基于局部视图构建逻辑通信环、分层拓扑或中继路径，将链路质量和节点处理能力纳入边权重；当节点加入、退出或链路异常时，仅在受影响的局部区段执行重构，并通过版本号、握手和短暂的新旧路径并存保持状态一致。研究将通过仿真和实际混合设备测试比较中心式全局调度与局部去中心化机制在控制开销、任务连续性和恢复时间方面的差异。",
        body_exemplar,
    )
    add_para(doc, elements, "3.2.4  多智能体强化学习动态适应方案", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "将调度过程建模为部分可观测的协作序贯决策问题，先利用统一模型和启发式策略构造稳定基线，再在小规模、低维状态空间中进行策略预训练，逐步扩展到更多节点和更复杂网络。算法选择不预先固定，按照任务规模和通信条件比较 MAPPO、QMIX、去中心化 Actor-Critic 以及约束强化学习等路线；训练阶段可以采用集中式信息，执行阶段仅使用局部观测。针对奖励稀疏、状态空间大和策略不稳定问题，引入课程学习、动作屏蔽、离线轨迹初始化和安全回退策略，确保学习策略失效时系统仍能退回可解释的启发式调度。",
        body_exemplar,
    )
    add_para(doc, elements, "3.2.5  原型系统与多场景实验方案", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "系统采用控制面、数据面和模型面分离的设计：控制面负责设备发现、任务编排、状态同步和策略下发；数据面负责 P2P KV 或中间数据传输；模型面通过后端接口适配 CUDA、ROCm、MPS 及后续可能接入的其他加速器。实验按“正确性—性能—扩展—鲁棒性—迁移性”五个维度开展：先与单设备和同构基线进行数值一致性对照，再测量不同设备组合、序列长度、网络条件和并发度下的性能，最后注入节点退出、链路抖动和负载变化，验证动态适应和容错机制。",
        body_exemplar,
    )

    add_para(doc, elements, "4. 预 期 达 到 的 目 标", heading_exemplar)
    add_para(
        doc,
        elements,
        "（1）理论目标：建立设备—网络—任务统一的异构合作模型，给出面向典型 AI 任务的可行性判定方法、资源约束表达和效率边界分析，形成能够迁移到推理、训练及其他分布式工作负载的建模框架。",
        body_exemplar,
    )
    add_para(
        doc,
        elements,
        "（2）方法目标：形成容量感知的非均匀任务划分、计算—通信联合调度、去中心化拓扑组织、局部容错重构和多智能体强化学习动态适应等关键方法，并明确各方法在设备差异、网络质量和任务结构变化下的适用范围。",
        body_exemplar,
    )
    add_para(
        doc,
        elements,
        "（3）系统目标：实现可插拔的异构协同原型系统，以 hetero-cp-ringattn 为首个实现载体，支持至少一种长上下文推理任务的端到端异构协同，并为分布式训练、强化学习工作流和其他任务保留统一接口。",
        body_exemplar,
    )
    add_para(
        doc,
        elements,
        "（4）实验目标：在多代、多厂商或不同设备能力的混合环境中，对比单设备、同构并行、静态异构调度和动态协同方法，系统评估完成时间、吞吐、显存、通信、能耗、扩展性和故障恢复能力；不预设单一固定加速倍数，而是报告不同条件下的收益区间和失效边界。",
        body_exemplar,
    )
    add_para(
        doc,
        elements,
        "（5）成果目标：围绕统一可行性建模、异构协同调度、去中心化容错和动态强化学习适配形成可拆分的论文与专利成果，持续维护并开源可复用的系统代码和实验基准。",
        body_exemplar,
    )

    add_para(doc, elements, "5. 已 完 成 的 研 究 工 作 与 进 度 安 排", heading_exemplar)
    add_para(doc, elements, "5.1  已完成的研究工作和取得的研究成果", subheading_exemplar)
    add_para(doc, elements, "5.1.1  已完成的研究工作", subsubheading_exemplar)
    add_para(doc, elements, "5.1.1.1  hetero-cp-ringattn 原型系统开发与验证", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "已完成 hetero-cp-ringattn 基础框架和异构 Ring Attention 通信原语的开发，形成设备性能探测、容量感知数据切片、KV ring 传输和可插拔传输后端等核心模块，并在小规模混合设备环境中完成正确性和性能测试。相关工作为后续研究统一设备画像、计算—通信联合调度和多平台后端适配提供了可运行的系统基础。",
        body_exemplar,
    )
    add_para(doc, elements, "5.1.1.2  异构分布式训练与推理系统调研", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "已完成对异构训练、异构推理、上下文并行、Prefill/Decode 分离和去中心化强化学习系统的调研，重点分析非对称并行、异构请求调度、P2P 协同和多智能体强化学习的设计思路，初步明确现有工作在单任务异构协同、局部拓扑适配和统一可行性分析方面的研究空间。",
        body_exemplar,
    )
    add_para(doc, elements, "5.1.1.3  异构设备画像与去中心化协同探索", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "已围绕设备算力、显存容量、链路带宽和通信延迟开展初步测量，完成动态环构建、非均匀切片和局部拓扑变化的仿真探索，并保留了跨 CUDA、ROCm 与 MPS 后端协同验证所需的接口设计。后续将把这些工程结果统一纳入可行性模型和动态适应实验。",
        body_exemplar,
    )
    add_para(doc, elements, "5.1.2  已取得的研究成果", subsubheading_exemplar)
    add_para(
        doc,
        elements,
        "（1）形成 hetero-cp-ringattn 开源代码库和异构集合通信原型；（2）完成异构硬件合作、上下文并行、分布式训练和去中心化强化学习方向的文献与系统调研；（3）搭建异构协同仿真与实验环境，具备开展设备画像、通信性能、动态重构和端到端推理验证的条件；（4）围绕异构集合通信和动态协同形成阶段性技术报告，并将根据后续实验结果拆分论文和专利选题。",
        body_exemplar,
    )
    add_para(doc, elements, "5.2  进度安排", subheading_exemplar)
    add_table(
        doc,
        elements,
        ["时间", "阶段目标", "主要产出"],
        [
            ["2026年9月—2026年12月", "完善设备与链路画像，建立统一任务—资源—网络模型；完成异构 Ring Attention 的正确性和基线测试。", "可行性模型初稿；设备画像数据集；阶段性实验报告。"],
            ["2027年1月—2027年4月", "研究容量感知任务划分、计算—通信联合调度和 P2P 拓扑组织；完善局部重构与可靠传输机制。", "异构协同调度原型；通信与拓扑模块；阶段性论文/专利素材。"],
            ["2027年5月—2027年8月", "研究多智能体强化学习动态适应，比较不同算法和启发式回退策略；在长上下文推理场景完成端到端验证。", "动态自适应模块；Prefill/Ring Attention 实验结果；论文初稿。"],
            ["2027年9月—2027年12月", "根据前期结果选择分布式训练、强化学习工作流或去中心化算力协作作为扩展验证场景，完成跨场景分析。", "扩展场景实验；适用边界分析；论文/专利投稿。"],
            ["2028年1月—2028年3月", "整合理论、算法和系统，完成综合实验、毕业论文撰写与答辩。", "开源代码与实验基准；毕业论文；答辩材料。"],
        ],
        style=doc.tables[0].style if doc.tables else None,
    )

    for element in elements:
        section6._p.addprevious(element)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
