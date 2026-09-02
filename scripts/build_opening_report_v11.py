from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path('/Users/stark_sim/Desktop/硕士课题/开题报告/开题报告_新版10_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx')
OUT = SRC.with_name('开题报告_新版11_沈达_面向大模型推理与强化学习后训练的异构算力负载承接研究.docx')


def set_text(para, text):
    ppr = deepcopy(para._p.pPr) if para._p.pPr is not None else None
    rpr = deepcopy(para.runs[0]._r.rPr) if para.runs and para.runs[0]._r.rPr is not None else None
    para.clear()
    run = para.add_run(text)
    if ppr is not None:
        para._p.insert(0, ppr)
    if rpr is not None:
        run._r.insert(0, rpr)


def remove_para(para):
    para._element.getparent().remove(para._element)


def first(doc, pred):
    return next(p for p in doc.paragraphs if pred(p.text))


def insert_after(para, text):
    new_p = deepcopy(para._p)
    for child in list(new_p):
        if child.tag != qn('w:pPr'):
            new_p.remove(child)
    para._p.addnext(new_p)
    new_para = next(p for p in para._parent.paragraphs if p._p is new_p)
    set_text(new_para, text)
    return new_para


def insert_after_with_template(para, text, template):
    new_p = deepcopy(template._p)
    para._p.addnext(new_p)
    new_para = next(p for p in para._parent.paragraphs if p._p is new_p)
    set_text(new_para, text)
    return new_para


def replace_all(doc, old, new):
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))


def update_toc(doc):
    replacements = {
        '2.1.1  多任务异构调度': '2.1.1  异构推理服务的请求与阶段承接',
        '2.1.2  单任务异构协同': '2.1.2  长上下文Prefill的任务内异构协同',
        '2.1.3  去中心化环境下的异构协同': '2.1.3  LLM-RL后训练的异构阶段承接与弹性执行',
        '2.1.4  去中心化与弱连接协作的适用边界': '2.1.4  去中心化与弱连接协作的适用边界',
        '2.2.1  异构计算负载均衡与调度': '2.2.1  异构推理系统与资源编排',
        '2.2.2  异构推理系统优化': '2.2.2  LLM-RL工作流与异构资源支持',
        '3. 主要研究内容及研究方案': '3. 主要研究内容及实施方案',
        '3.1  主要研究内容': '3.1  研究内容',
        '3.2  研究方案': '3.2  实施方案',
    }
    for p in doc.part._element.xpath('.//w:sdt//w:p'):
        text = ''.join(t.text or '' for t in p.iter(qn('w:t')))
        for old, new in replacements.items():
            if old in text:
                ts = list(p.iter(qn('w:t')))
                # TOC entries commonly split the heading across two runs and
                # keep the page number in the final run. Replace only heading
                # runs so page fields and their formatting remain untouched.
                page_run = ts[-1] if ts and (ts[-1].text or '').strip().isdigit() else None
                heading_ts = ts[:-1] if page_run is not None else ts
                if heading_ts:
                    heading_ts[0].text = new
                    for t in heading_ts[1:]:
                        t.text = ''
                break
    # The source TOC has no 2.1.4 row. Clone the existing 2.1.3 row so the
    # new boundary subsection appears with the same tabs and page-number run.
    toc_ps = doc.part._element.xpath('.//w:sdt//w:p')
    marker = None
    for p in toc_ps:
        text = ''.join(t.text or '' for t in p.iter(qn('w:t')))
        if text.startswith('2.1.3  LLM-RL后训练的异构阶段承接与弹性执行'):
            marker = p
            break
    if marker is not None and not any(
        ''.join(t.text or '' for t in p.iter(qn('w:t'))).startswith('2.1.4  去中心化与弱连接协作的适用边界')
        for p in toc_ps
    ):
        clone = deepcopy(marker)
        ts = list(clone.iter(qn('w:t')))
        page_run = ts[-1] if ts and (ts[-1].text or '').strip().isdigit() else None
        heading_ts = ts[:-1] if page_run is not None else ts
        if heading_ts:
            heading_ts[0].text = '2.1.4  去中心化与弱连接协作的适用边界'
            for t in heading_ts[1:]:
                t.text = ''
        if page_run is not None:
            page_run.text = '4'
        marker.addnext(clone)


def main():
    doc = Document(SRC)

    # Add a wider entry point before introducing the specific workload scope.
    h11 = first(doc, lambda t: t.strip() == '1.1  课题的来源')
    body = first(doc, lambda t: t.startswith('近年来，大语言模型推理和LLM-RL后训练'))
    anchor = h11
    for text in [
        '人工智能技术正在从以模型训练为中心的研发阶段，逐步进入以持续服务、行业应用和模型能力迭代为特征的工程化阶段。大语言模型、视觉语言模型和面向科学计算的基础模型不断扩大参数规模与上下文窗口，模型的使用方式也从离线批处理扩展到在线问答、代码生成、智能体执行和多轮交互。对于系统而言，这意味着算力需求不再只表现为一次性的训练峰值，而是同时表现为长期运行的推理请求、周期性的模型评测以及持续进行的后训练任务；推理服务的阶段拆分和资源池化已经成为公开系统持续演进的重要方向[1-5]。',
        '算力基础设施的组织方式也随之发生变化。一方面，头部云服务商和大型机构可以围绕少数型号建设高带宽、强互联的同构集群；另一方面，实际环境往往由不同采购批次、不同厂商和不同代际的设备共同组成，实验室、企业私有集群以及跨机构算力联盟尤其如此。开放互连标准、专用AI网络和国产超节点的发展，也使不同硬件在同一基础设施中长期共存成为需要正视的工程事实[21-24]。设备之间的显存、算力、驱动栈和互联条件并不一致，部分设备适合高并发推理，部分设备更适合批量生成或评估，另一些设备则只能在通信较少的阶段发挥作用。如何把这些设备组织成有边界、可回退的有效资源，已经成为模型系统落地时无法回避的问题。',
        '从负载结构看，异构资源并非只能通过“把同一个任务平均切成若干份”来使用。在线推理可以按请求、Prefill/Decode阶段和KV缓存状态进行组织；LLM-RL后训练可以按rollout、奖励计算、评估、数据处理和策略更新等阶段进行组织；只有在长上下文造成单设备显存不足时，才有必要进一步考虑一次注意力计算内部的上下文切分。因此，异构承接首先是负载识别和协同粒度选择问题，其次才是并行算法和通信协议问题。',
    ]:
        anchor = insert_after_with_template(anchor, text, body)

    # Expand the significance section from system trends to the concrete gap.
    h12 = first(doc, lambda t: t.strip() == '1.2  课题研究的背景和意义')
    body12 = first(doc, lambda t: t.startswith('与同构集群相比，异构集群'))
    anchor = h12
    for text in [
        '模型规模和服务形态的变化，使“算力是否足够”逐渐转化为“算力是否能够被有效组织”。训练阶段通常依赖稳定的高带宽互联和较强的一致性，推理阶段则同时受到请求到达、上下文长度、输出长度、KV缓存和服务等级目标的影响；相关系统研究已经从单体推理逐步发展到Prefill/Decode分离、KV管理和多级资源编排[1-5]。后训练还处于训练与服务的交界处：它既需要生成大量样本，又需要周期性地进行奖励计算、评估和策略更新。不同阶段的通信强度、计算密度和容错要求并不相同，不能使用单一的并行方式覆盖所有环节。',
        '因此，行业中的系统演进大致形成了由粗到细的三层路径。第一层是资源池和请求级调度，将不同请求分配到合适的设备或副本；第二层是阶段级编排，把Prefill与Decode、rollout与策略更新等具有边界的阶段分开；第三层才是任务内并行，在单个请求或单次计算内部切分序列、激活或KV数据。前两层通常可以容忍较弱的设备互联，第三层则必须把通信代价和同步频率纳入正确性与性能模型。',
        '本课题选择第三层中的一个窄而明确的推理问题，以及第二层中的一个后训练问题进行研究，原因在于二者既能体现异构设备的实际承接价值，又具有不同的风险边界。长上下文Prefill的核心困难是单设备显存和一次计算内的K/V交换；LLM-RL阶段承接的核心困难是阶段能力不匹配、异步版本滞后和节点变化。把两者并列讨论，可以避免将“请求级路由有效”不恰当地外推为“任意任务内并行有效”，也避免把训练侧的弱同步结论直接套用到推理侧的紧耦合通信。',
    ]:
        anchor = insert_after_with_template(anchor, text, body12)

    # Give Chapter 2 a reader-oriented map before the detailed survey.
    h2 = first(doc, lambda t: t.strip() == '2. 国内外在该方向的研究现状及分析')
    body2 = first(doc, lambda t: t.startswith('国外推理系统首先在请求和阶段层利用异构性'))
    anchor = h2
    for text in [
        '大模型系统研究并不是从“异构卡共同执行一个算子”起步的。随着模型规模和请求数量增加，研究者先解决了模型如何部署、请求如何排队以及推理阶段如何拆分等基础问题，随后才逐步讨论KV缓存迁移、长上下文并行、动态资源编排和跨地域协作。理解这一演进顺序，有助于区分已有系统能力与本课题拟增加的机制，也能避免把成熟的服务编排能力误写成本课题创新。',
        '本章按“服务组织—任务内协同—后训练工作流—弱连接协作”的顺序梳理国外研究，再从国内公开工作和工程实践进行补充。前两类工作主要对应推理线，后两类工作主要对应LLM-RL线；最后从支持证据、反对证据和适用边界三个方面归纳现有结论。这样的整理不预设异构合作一定有效，而是先说明哪些问题已经有成熟答案，哪些问题仍缺乏可复现实验。',
    ]:
        anchor = insert_after_with_template(anchor, text, body2)

    # Add a short transition between overseas and domestic evidence.
    h22 = first(doc, lambda t: t.strip() == '2.2  国内研究现状')
    body22 = first(doc, lambda t: t.startswith('国内相关工作在异构推理的资源编排'))
    anchor = h22
    for text in [
        '从国内外研究的整体演进看，公开成果已经覆盖了异构资源编排、KV管理、长上下文并行和RL工作流等多个环节，但不同成果通常围绕各自的目标系统建立，评价指标和硬件假设并不完全一致。国内研究更加重视面向具体集群的部署效率、内存层次和资源利用，国外研究则形成了较多面向开放集群、异步工作流和跨地域协作的公开系统。将这些成果放在同一框架下比较时，应关注其协同粒度、通信频率、状态一致性和失败处理方式，而不能只比较设备数量或峰值吞吐。',
    ]:
        anchor = insert_after_with_template(anchor, text, body22)

    update_toc(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    main()
